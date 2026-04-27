from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from backend.contracts import AuditReport, Issue, ObjectRef, Severity
from backend.core import _build_rule_hit_notes, _finalize_issues, _infer_drawing_type
from backend.model.types import DrawingSelection
from backend.reporting.docx_report import write_docx_report
from backend.reporting.markdown import render_markdown_report


class RuleHitNotesTests(unittest.TestCase):
    def test_single_line_drawings_escalate_third_batch_rule_to_error(self) -> None:
        selection = DrawingSelection(
            drawing_class='supported_electrical',
            reason='matched_supported_keyword:380v',
            eligible_for_electrical=True,
        )
        issues = _finalize_issues(
            [
                Issue(
                    rule_id='electrical.incoming_transformer_busbar_direction',
                    severity=Severity.WARNING,
                    message='x',
                )
            ],
            selection=selection,
            input_path=Path('380V.dwg'),
        )

        self.assertEqual(issues[0].severity, Severity.ERROR)

    def test_electrical_schematic_drawings_keep_third_batch_rule_as_warning(self) -> None:
        selection = DrawingSelection(
            drawing_class='supported_electrical',
            reason='matched_supported_keyword',
            eligible_for_electrical=True,
        )
        issues = _finalize_issues(
            [
                Issue(
                    rule_id='electrical.tie_busbar_segment_consistency',
                    severity=Severity.WARNING,
                    message='x',
                )
            ],
            selection=selection,
            input_path=Path('低压综合配电箱电气图.dxf'),
        )

        self.assertEqual(issues[0].severity, Severity.WARNING)

    def test_rule_hit_notes_include_drawing_type_and_reason(self) -> None:
        selection = DrawingSelection(
            drawing_class='supported_electrical',
            reason='matched_supported_keyword',
            eligible_for_electrical=True,
        )
        notes = _build_rule_hit_notes(
            (
                Issue(
                    rule_id='electrical.tie_busbar_segment_consistency',
                    severity=Severity.WARNING,
                    message='x',
                ),
            ),
            selection=selection,
            input_path=Path('主接线.dxf'),
        )

        self.assertEqual(notes[0]['drawing_type'], 'single_line')
        self.assertIn('error', str(notes[0]['grading_reason']).lower())

    def test_markdown_report_renders_rule_hit_notes_section(self) -> None:
        report = AuditReport(
            created_at='2026-03-24T10:00:00+08:00',
            input_path='D:/tmp/主接线.dxf',
            input_sha256='abc',
            parser='dxf_ascii_v1',
            rule_version='rules_v1',
            issues=(
                Issue(
                    rule_id='electrical.incoming_transformer_busbar_direction',
                    severity=Severity.ERROR,
                    message='进线柜方向关系异常',
                    refs=(
                        ObjectRef(
                            kind='terminal',
                            id='t1',
                            extra={'x': 12.5, 'y': 8.0},
                        ),
                        ObjectRef(
                            kind='terminal',
                            id='t2',
                            extra={'x': 24.0, 'y': 8.0},
                        ),
                    ),
                ),
            ),
            summary={
                'classification': {'drawing_type': 'single_line'},
                'rule_hit_notes': [
                    {
                        'rule_id': 'electrical.incoming_transformer_busbar_direction',
                        'title': '进线柜-变压器-母线方向一致性',
                        'count': 1,
                        'severity': 'error',
                        'drawing_type_label': '单线/一次系统图',
                        'meaning': '检查一次侧与母线侧是否相对。',
                        'grading_reason': '该类图纸直接表达一次接线与主供电关系，命中后按 error 处理。',
                    }
                ],
            },
        )

        rendered = render_markdown_report(report)

        self.assertIn('## Rule Hit Notes', rendered)
        self.assertIn('单线/一次系统图', rendered)
        self.assertIn('grading_reason', rendered)

    def test_markdown_report_renders_formal_issue_fields(self) -> None:
        report = AuditReport(
            created_at='2026-03-24T10:00:00+08:00',
            input_path='D:/tmp/主接线.dxf',
            input_sha256='abc',
            parser='dxf_ascii_v1',
            rule_version='rules_v1',
            issues=(
                Issue(
                    rule_id='electrical.incoming_transformer_busbar_direction',
                    severity=Severity.ERROR,
                    message='进线柜方向关系异常',
                    refs=(
                        ObjectRef(
                            kind='terminal',
                            id='t1',
                            extra={'x': 12.5, 'y': 8.0},
                        ),
                        ObjectRef(
                            kind='terminal',
                            id='t2',
                            extra={'x': 24.0, 'y': 8.0},
                        ),
                    ),
                ),
            ),
            summary={
                'classification': {
                    'drawing_type': 'single_line',
                    'drawing_type_label': '单线/一次系统图',
                },
            },
        )

        rendered = render_markdown_report(report)

        self.assertIn('article_clause_mapping: SF-EL-012', rendered)
        self.assertIn('remediation:', rendered)
        self.assertIn('risk_level: high', rendered)
        self.assertIn('confidence: high', rendered)

    def test_docx_report_renders_formal_issue_fields(self) -> None:
        from docx import Document

        report = AuditReport(
            created_at='2026-03-24T10:00:00+08:00',
            input_path='D:/tmp/主接线.dxf',
            input_sha256='abc',
            parser='dxf_ascii_v1',
            rule_version='rules_v1',
            issues=(
                Issue(
                    rule_id='electrical.tie_busbar_segment_consistency',
                    severity=Severity.WARNING,
                    message='联络柜母线分段异常',
                ),
            ),
            summary={
                'classification': {
                    'drawing_type': 'single_line',
                    'drawing_type_label': '单线/一次系统图',
                },
            },
        )

        with tempfile.TemporaryDirectory() as td:
            out_path = Path(td) / 'report.docx'
            write_docx_report(report, out_path)

            doc = Document(str(out_path))
            text_parts = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

        rendered = '\n'.join(text_parts)
        self.assertIn('article_clause_mapping', rendered)
        self.assertIn('SF-EL-013 联络柜两侧母线分段独立性', rendered)
        self.assertIn('risk_level', rendered)
        self.assertIn('confidence', rendered)
        self.assertIn('remediation', rendered)

    def test_infer_drawing_type_marks_electrical_layout_as_layout(self) -> None:
        selection = DrawingSelection(
            drawing_class='geometry_only',
            reason='matched_parent_dir',
            eligible_for_electrical=False,
        )
        self.assertEqual(_infer_drawing_type(Path('低压综合配电箱布置加工图.dwg'), selection), 'layout_or_installation')


if __name__ == '__main__':
    unittest.main()
