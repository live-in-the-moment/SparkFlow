from __future__ import annotations

import contextlib
import hashlib
import io
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from sparkflow.__main__ import main
from sparkflow.reporting.rectification_checklist import write_rectification_checklist


class RectificationChecklistTests(unittest.TestCase):
    def test_rectification_checklist_writes_failed_drawings_only_markdown_docx_and_json(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            run_dir = _build_run_dir(root)

            artifacts = write_rectification_checklist(run_dir)

            markdown = artifacts.markdown_path.read_text(encoding='utf-8')
            self.assertIn('配电部分CAD 严格审图整改清单', markdown)
            self.assertIn('纳入整改清单的失败电气图纸：5', markdown)
            self.assertIn('整改项总数：59', markdown)
            self.assertIn('低压综合配电箱/DP-2/低压综合配电箱DP-2图纸改.dwg', markdown)
            self.assertIn('优先复核坐标 (1337.693, 3045.297) 附近对象', markdown)
            self.assertIn('files/低压综合配电箱/DP-2/低压综合配电箱DP-2图纸改__', markdown)
            self.assertNotIn('电缆分支箱/电缆分支箱DF-2/一次系统图.dwg', markdown)
            self.assertNotIn('低压开关柜DK-1/平面.dwg', markdown)
            self.assertNotIn('????', markdown)

            payload = json.loads(artifacts.json_path.read_text(encoding='utf-8'))
            self.assertEqual(payload['counts']['included_failed_drawings'], 5)
            self.assertEqual(payload['counts']['total_issue_count'], 59)
            self.assertEqual(len(payload['drawings']), 5)
            self.assertEqual(sum(len(drawing['issues']) for drawing in payload['drawings']), 59)

            extracted = _extract_docx_text(artifacts.docx_path)
            self.assertIn('配电部分CAD 严格审图整改清单', extracted)
            self.assertIn('整改项总数：59', extracted)
            self.assertIn('低压开关柜DK-1/主接线.dwg', extracted)
            self.assertIn('优先复核坐标', extracted)
            self.assertNotIn('????', extracted)

    def test_rectification_checklist_command_writes_default_outputs(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            run_dir = _build_run_dir(root)

            stdout = io.StringIO()
            stderr = io.StringIO()
            with contextlib.redirect_stdout(stdout), contextlib.redirect_stderr(stderr):
                exit_code = main(['rectification-checklist', str(run_dir)])

            self.assertEqual(exit_code, 0)
            lines = stdout.getvalue().splitlines()
            self.assertEqual(len(lines), 3)
            self.assertTrue(Path(lines[0]).exists())
            self.assertTrue(Path(lines[1]).exists())
            self.assertTrue(Path(lines[2]).exists())
            self.assertEqual(stderr.getvalue(), '')


def _build_run_dir(root: Path) -> Path:
    run_dir = root / '20260325T120000Z'
    run_dir.mkdir(parents=True, exist_ok=True)
    (run_dir / 'dataset_summary.json').write_text(
        json.dumps(
            {
                'created_at': '2026-03-25T12:00:00+08:00',
                'dataset_dir': 'D:\\code\\project\\moment\\SparkFlow\\image\\国家电网公司380220V配电网工程典型设计（2018年版）_1772430671059\\配电部分CAD',
                'rule_version': 'stategrid_peidian_strict_v1',
                'counts': {'passed': 1, 'failed': 5, 'skipped': 2, 'unprocessed': 0},
                'selection_counts': {'supported_electrical': 6, 'geometry_only': 1, 'unsupported': 1},
                'issues_by_rule': {'wire.floating_endpoints': 59},
                'failures': [],
                'unprocessed': [],
                'timing': {'elapsed_sec': 12.34, 'avg_file_sec': 2.056, 'file_count': 8},
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding='utf-8',
    )
    selection_rows = [
        {
            'rel_path': '低压开关柜DK-1/380V.dwg',
            'drawing_class': 'supported_electrical',
            'reason': 'matched_supported_keyword:380v',
            'eligible_for_electrical': True,
            'status': 'failed',
            'elapsed_sec': 5.0,
        },
        {
            'rel_path': '低压开关柜DK-1/主接线.dwg',
            'drawing_class': 'supported_electrical',
            'reason': 'matched_supported_keyword:主接线',
            'eligible_for_electrical': True,
            'status': 'failed',
            'elapsed_sec': 5.0,
        },
        {
            'rel_path': '低压综合配电箱/DP-2/低压综合配电箱DP-2图纸改.dwg',
            'drawing_class': 'supported_electrical',
            'reason': 'text_electrical_match:配电箱|变压器|避雷器',
            'eligible_for_electrical': True,
            'status': 'failed',
            'elapsed_sec': 5.0,
        },
        {
            'rel_path': '低压综合配电箱/DP-3/低压综合配电箱电气图.dwg',
            'drawing_class': 'supported_electrical',
            'reason': 'matched_supported_keyword:电气图',
            'eligible_for_electrical': True,
            'status': 'failed',
            'elapsed_sec': 5.0,
        },
        {
            'rel_path': '低压综合配电箱/DP-4/电气图-DP-4.dwg',
            'drawing_class': 'supported_electrical',
            'reason': 'matched_supported_keyword:电气图',
            'eligible_for_electrical': True,
            'status': 'failed',
            'elapsed_sec': 5.0,
        },
        {
            'rel_path': '电缆分支箱/电缆分支箱DF-2/一次系统图.dwg',
            'drawing_class': 'supported_electrical',
            'reason': 'matched_supported_keyword:一次系统图',
            'eligible_for_electrical': True,
            'status': 'passed',
            'elapsed_sec': 5.0,
        },
        {
            'rel_path': '低压开关柜DK-1/平面.dwg',
            'drawing_class': 'geometry_only',
            'reason': 'matched_geometry_keyword:平面',
            'eligible_for_electrical': False,
            'status': 'skipped',
            'elapsed_sec': 0.2,
        },
        {
            'rel_path': '电缆分支箱/低压典设电缆分支箱图纸-DF-1,DF-3,DF-4.dwg',
            'drawing_class': 'unsupported',
            'reason': 'no_supported_keyword_match',
            'eligible_for_electrical': False,
            'status': 'skipped',
            'elapsed_sec': 0.2,
        },
    ]
    (run_dir / 'dataset_selection.json').write_text(
        json.dumps(selection_rows, ensure_ascii=False, indent=2),
        encoding='utf-8',
    )
    _write_file_report(
        run_dir,
        rel_path='低压开关柜DK-1/380V.dwg',
        drawing_class='supported_electrical',
        drawing_type='single_line',
        drawing_type_label='单线/一次系统图',
        status='failed',
        issues=[
            _issue(17467.38205166408, -23445.76818192429),
        ],
    )
    _write_file_report(
        run_dir,
        rel_path='低压开关柜DK-1/主接线.dwg',
        drawing_class='supported_electrical',
        drawing_type='single_line',
        drawing_type_label='单线/一次系统图',
        status='failed',
        issues=[_issue(17500.0 + index, -23500.0 - index) for index in range(11)],
    )
    dp2_coords = [
        (1337.693096296129, 3045.297044305588),
        (1341.593140279774, 3045.297044305588),
        (1327.997718039451, 3068.267942772289),
        (1383.15125778354, 3068.267942772289),
        (1340.954121332283, 3141.238806568485),
        (1340.704441765689, 3178.725484724886),
        (1338.549493699729, 3173.632018142506),
        (1338.333327595518, 3141.238562474681),
        (1341.8145308853, 3065.780798201651),
        (1914.835242254842, 3171.980766465949),
        (1914.835242254842, 3143.981027097833),
        (1355.712245524196, 2871.262352132086),
        (1356.962245524196, 2872.512352132086),
        (1355.63978061549, 2872.512352132086),
        (1356.88978061549, 2871.262352132086),
        (1407.079135656738, 2896.887404372985),
        (1407.385890015433, 2942.85592227436),
        (1410.456152027388, 2944.350172523012),
        (1410.784593282067, 2895.083984325243),
        (1427.683758730916, 2936.3304248085),
        (1427.743785755508, 2945.325743944453),
        (1430.798230510935, 2944.256153769394),
        (1430.851068704072, 2936.3304248085),
        (1328.421794125702, 2870.7525337163593),
        (1327.07779658822, 2872.512352132086),
        (1382.609052056539, 2872.512352132086),
        (1338.840279413305, 2903.07854034457),
        (1333.8260238772164, 2898.916970943148),
        (1401.7878817564826, 2930.725259439823),
        (1914.835242254842, 2961.291322349853),
        (1914.835242254842, 2919.291713297677),
    ]
    _write_file_report(
        run_dir,
        rel_path='低压综合配电箱/DP-2/低压综合配电箱DP-2图纸改.dwg',
        drawing_class='supported_electrical',
        drawing_type='general_supported_electrical',
        drawing_type_label='一般电气审图图纸',
        status='failed',
        issues=[_issue(x, y) for x, y in dp2_coords],
    )
    _write_file_report(
        run_dir,
        rel_path='低压综合配电箱/DP-3/低压综合配电箱电气图.dwg',
        drawing_class='supported_electrical',
        drawing_type='electrical_schematic',
        drawing_type_label='电气原理图',
        status='failed',
        issues=[_issue(1100.0 + index, 2200.0 + index) for index in range(5)],
    )
    _write_file_report(
        run_dir,
        rel_path='低压综合配电箱/DP-4/电气图-DP-4.dwg',
        drawing_class='supported_electrical',
        drawing_type='electrical_schematic',
        drawing_type_label='电气原理图',
        status='failed',
        issues=[_issue(2100.0 + index, 3200.0 + index) for index in range(11)],
    )
    _write_file_report(
        run_dir,
        rel_path='电缆分支箱/电缆分支箱DF-2/一次系统图.dwg',
        drawing_class='supported_electrical',
        drawing_type='single_line',
        drawing_type_label='单线/一次系统图',
        status='passed',
        issues=[],
    )
    _write_file_report(
        run_dir,
        rel_path='低压开关柜DK-1/平面.dwg',
        drawing_class='geometry_only',
        drawing_type='layout_or_installation',
        drawing_type_label='布置/安装类图纸',
        status='skipped',
        issues=[],
    )
    _write_file_report(
        run_dir,
        rel_path='电缆分支箱/低压典设电缆分支箱图纸-DF-1,DF-3,DF-4.dwg',
        drawing_class='unsupported',
        drawing_type='other',
        drawing_type_label='暂不支持图纸',
        status='skipped',
        issues=[],
    )
    return run_dir


def _write_file_report(
    run_dir: Path,
    *,
    rel_path: str,
    drawing_class: str,
    drawing_type: str,
    drawing_type_label: str,
    status: str,
    issues: list[dict[str, object]],
) -> None:
    file_dir = _dataset_file_out_dir(run_dir, rel_path)
    file_dir.mkdir(parents=True, exist_ok=True)
    report = {
        'created_at': '2026-03-25T12:00:00+08:00',
        'input_path': f'D:\\fake\\{rel_path}',
        'input_sha256': 'abc',
        'parser': 'stub',
        'rule_version': 'stategrid_peidian_strict_v1',
        'passed': status == 'passed',
        'issues': issues,
        'summary': {
            'classification': {
                'drawing_class': drawing_class,
                'reason': 'stub',
                'eligible_for_electrical': drawing_class == 'supported_electrical',
                'drawing_type': drawing_type,
                'drawing_type_label': drawing_type_label,
            }
        },
        'artifacts': {'report_docx': 'report.docx'},
    }
    (file_dir / 'report.json').write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding='utf-8')
    (file_dir / 'report.md').write_text(f'# {rel_path}\n', encoding='utf-8')
    Document().save(str(file_dir / 'report.docx'))


def _issue(x: float, y: float) -> dict[str, object]:
    return {
        'rule_id': 'wire.floating_endpoints',
        'severity': 'error',
        'message': '发现悬空线端点，未与其他线段或设备连接。',
        'refs': [{'kind': 'node', 'id': f'node:{x}', 'extra': {'x': x, 'y': y, 'degree': 1}}],
    }


def _dataset_file_out_dir(run_dir: Path, rel_path: str) -> Path:
    path = Path(*rel_path.split('/'))
    token = hashlib.sha1(rel_path.encode('utf-8')).hexdigest()[:8]
    return run_dir / 'files' / path.parent / f'{path.stem}__{token}'


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return '\n'.join(part for part in parts if part)


if __name__ == '__main__':
    unittest.main()
