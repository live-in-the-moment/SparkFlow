from __future__ import annotations

import contextlib
import hashlib
import io
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from backend.__main__ import main
from backend.reporting.dataset_report import write_dataset_audit_report


class DatasetReportTests(unittest.TestCase):
    def test_dataset_report_writes_chinese_markdown_and_docx(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            run_dir = _build_run_dir(root)
            ruleset_dir = _build_ruleset_dir(root / 'ruleset')

            artifacts = write_dataset_audit_report(run_dir, ruleset_dir=ruleset_dir)

            markdown = artifacts.markdown_path.read_text(encoding='utf-8')
            self.assertIn('配电部分CAD 严格审图最终报告', markdown)
            self.assertIn('严格规则：wire.floating_endpoints 已按 error 判定', markdown)
            self.assertIn('低压开关柜DK-1/主接线.dwg', markdown)
            self.assertIn('跳过的布置/几何图纸', markdown)
            self.assertIn('files/低压开关柜DK-1/主接线__', markdown)
            self.assertNotIn('????', markdown)

            extracted = _extract_docx_text(artifacts.docx_path)
            self.assertIn('配电部分CAD 严格审图最终报告', extracted)
            self.assertIn('低压开关柜DK-1/主接线.dwg', extracted)
            self.assertIn('严格规则：wire.floating_endpoints 已按 error 判定', extracted)
            self.assertNotIn('????', extracted)

    def test_dataset_report_command_writes_default_outputs(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            run_dir = _build_run_dir(root)
            ruleset_dir = _build_ruleset_dir(root / 'ruleset')

            stdout = io.StringIO()
            stderr = io.StringIO()
            with contextlib.redirect_stdout(stdout), contextlib.redirect_stderr(stderr):
                exit_code = main(['dataset-report', str(run_dir), '--ruleset', str(ruleset_dir)])

            self.assertEqual(exit_code, 0)
            lines = stdout.getvalue().splitlines()
            self.assertEqual(len(lines), 2)
            self.assertTrue(Path(lines[0]).exists())
            self.assertTrue(Path(lines[1]).exists())
            self.assertEqual(stderr.getvalue(), '')


def _build_run_dir(root: Path) -> Path:
    run_dir = root / '20260325T120000Z'
    run_dir.mkdir(parents=True, exist_ok=True)
    (run_dir / 'dataset_summary.json').write_text(
        json.dumps(
            {
                'created_at': '2026-03-25T12:00:00+08:00',
                'dataset_dir': 'D:\\path\\dataset_root\\配电部分CAD',
                'rule_version': 'stategrid_peidian_strict_v1',
                'counts': {'passed': 1, 'failed': 1, 'skipped': 1, 'unprocessed': 0},
                'selection_counts': {'supported_electrical': 2, 'geometry_only': 1, 'unsupported': 0},
                'issues_by_rule': {'wire.floating_endpoints': 3},
                'failures': [],
                'unprocessed': [],
                'timing': {'elapsed_sec': 12.34, 'avg_file_sec': 4.113, 'file_count': 3},
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding='utf-8',
    )
    selection_rows = [
        {
            'rel_path': '低压开关柜DK-1/主接线.dwg',
            'drawing_class': 'supported_electrical',
            'reason': 'matched_supported_keyword:主接线',
            'eligible_for_electrical': True,
            'status': 'failed',
            'elapsed_sec': 5.0,
        },
        {
            'rel_path': '电缆分支箱/电缆分支箱DF-2/一次系统图.dwg',
            'drawing_class': 'supported_electrical',
            'reason': 'matched_supported_keyword:一次系统图',
            'eligible_for_electrical': True,
            'status': 'passed',
            'elapsed_sec': 4.0,
        },
        {
            'rel_path': '低压开关柜DK-1/平面.dwg',
            'drawing_class': 'geometry_only',
            'reason': 'matched_geometry_keyword:平面',
            'eligible_for_electrical': False,
            'status': 'skipped',
            'elapsed_sec': 0.2,
        },
    ]
    (run_dir / 'dataset_selection.json').write_text(
        json.dumps(selection_rows, ensure_ascii=False, indent=2),
        encoding='utf-8',
    )

    _write_file_report(
        run_dir,
        rel_path='低压开关柜DK-1/主接线.dwg',
        drawing_class='supported_electrical',
        drawing_type_label='单线/一次系统图',
        status='failed',
        issues=[
            _issue(17467.4, -23445.8),
            _issue(17552.5, -23446.1),
            _issue(17603.1, -23473.0),
        ],
    )
    _write_file_report(
        run_dir,
        rel_path='电缆分支箱/电缆分支箱DF-2/一次系统图.dwg',
        drawing_class='supported_electrical',
        drawing_type_label='单线/一次系统图',
        status='passed',
        issues=[],
    )
    _write_file_report(
        run_dir,
        rel_path='低压开关柜DK-1/平面.dwg',
        drawing_class='geometry_only',
        drawing_type_label='布置/安装类图纸',
        status='skipped',
        issues=[],
    )
    return run_dir


def _write_file_report(
    run_dir: Path,
    *,
    rel_path: str,
    drawing_class: str,
    drawing_type_label: str,
    status: str,
    issues: list[dict[str, object]],
) -> None:
    file_dir = _dataset_file_out_dir(run_dir, rel_path)
    file_dir.mkdir(parents=True, exist_ok=True)
    report = {
        'created_at': '2026-03-25T12:00:00+08:00',
        'input_path': f'D:\\fake\\{rel_path}',
        'input_sha256': 'abc',
        'parser': 'stub',
        'rule_version': 'stategrid_peidian_strict_v1',
        'passed': status == 'passed',
        'issues': issues,
        'summary': {
            'classification': {
                'drawing_class': drawing_class,
                'reason': 'stub',
                'eligible_for_electrical': drawing_class == 'supported_electrical',
                'drawing_type': 'single_line' if drawing_class == 'supported_electrical' else 'layout_or_installation',
                'drawing_type_label': drawing_type_label,
            }
        },
        'artifacts': {'report_docx': 'report.docx'},
    }
    (file_dir / 'report.json').write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding='utf-8')
    (file_dir / 'report.md').write_text(f'# {rel_path}\n', encoding='utf-8')
    Document().save(str(file_dir / 'report.docx'))


def _build_ruleset_dir(ruleset_dir: Path) -> Path:
    ruleset_dir.mkdir(parents=True, exist_ok=True)
    (ruleset_dir / 'ruleset.json').write_text(
        json.dumps(
            {
                'version': 'stategrid_peidian_strict_v1',
                'rules': [
                    {
                        'rule_id': 'wire.floating_endpoints',
                        'enabled': True,
                        'severity': 'error',
                        'params': {'tol': 0.001},
                        'applies_to': ['single_line', 'electrical_schematic', 'general_supported_electrical'],
                    }
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding='utf-8',
    )
    return ruleset_dir


def _issue(x: float, y: float) -> dict[str, object]:
    return {
        'rule_id': 'wire.floating_endpoints',
        'severity': 'error',
        'message': '发现悬空线端点，未与其他线段或设备连接。',
        'refs': [{'kind': 'node', 'id': f'node:{x}', 'extra': {'x': x, 'y': y}}],
    }


def _dataset_file_out_dir(run_dir: Path, rel_path: str) -> Path:
    path = Path(*rel_path.split('/'))
    token = hashlib.sha1(rel_path.encode('utf-8')).hexdigest()[:8]
    return run_dir / 'files' / path.parent / f'{path.stem}__{token}'


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return '\n'.join(part for part in parts if part)


if __name__ == '__main__':
    unittest.main()
