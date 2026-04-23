from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from openpyxl import Workbook

from sparkflow.model.types import ProjectDocumentContext
from sparkflow.project_docs import build_project_document_context, project_document_mentions, project_document_note_exists


class ProjectDocumentContextTests(unittest.TestCase):
    def test_build_project_document_context_extracts_counts_from_docx_xlsx_and_xls(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            project_root = Path(td) / "工程A"
            drawing_path = project_root / "附件3 施工图" / "一次系统图.dxf"
            drawing_path.parent.mkdir(parents=True, exist_ok=True)
            drawing_path.write_text(_DXF_TEXT_ONLY, encoding="utf-8")

            _write_docx(
                project_root / "附件1 施工图设计说明书_工程A.docx",
                rows=(
                    ("低压配电网", "配变数量（台）", "配变数量（台）", "2", "2", "0"),
                    ("智能化", "配电智能网关", "配电智能网关", "2", "2", "0"),
                    ("配网自动化", "DTU（台）", "DTU（台）", "1", "1", "0"),
                ),
                extra_paragraphs=("二次柜工程量说明详见设备选型表。",),
            )
            _write_xlsx(
                project_root / "附件4 主要设备材料清册.xlsx",
                rows=(
                    ("设备名称", "数量", "单位"),
                    ("综合测控通信单元", "1", "台"),
                    ("二次柜", "2", "台"),
                ),
            )
            tower_xls = project_root / "附件3 施工图" / "10kV及以下杆（塔）明细表-工程A.xls"
            tower_xls.write_bytes(b"stub")

            with patch(
                "sparkflow.project_docs._read_xls_rows_via_excel",
                return_value=[
                    ["名称", "数量", "单位"],
                    ["电杆", "3", "根"],
                ],
            ):
                context = build_project_document_context(drawing_path)

            assert context is not None
            self.assertEqual(context.expected_counts["distribution_transformer"], 2.0)
            self.assertEqual(context.expected_counts["smart_gateway"], 2.0)
            self.assertEqual(context.expected_counts["dtu"], 1.0)
            self.assertEqual(context.expected_counts["measurement_comm_unit"], 1.0)
            self.assertEqual(context.expected_counts["secondary_cabinet"], 2.0)
            self.assertEqual(context.expected_counts["pole_tower"], 3.0)
            self.assertTrue(project_document_note_exists(context, "secondary_cabinet"))
            self.assertEqual(len(context.sources), 3)
            self.assertEqual(context.errors, ())

    def test_project_document_mentions_zero_count_matrix(self) -> None:
        zero_only = ProjectDocumentContext(project_root="D:/project", expected_counts={"secondary_cabinet": 0.0})
        self.assertFalse(project_document_mentions(zero_only, "secondary_cabinet"))

        positive = ProjectDocumentContext(project_root="D:/project", expected_counts={"secondary_cabinet": 2.0})
        self.assertTrue(project_document_mentions(positive, "secondary_cabinet"))

        text_only = ProjectDocumentContext(
            project_root="D:/project",
            expected_counts={"secondary_cabinet": 0.0},
            text_snippets=("二次柜安装位置详见平面布置图。",),
        )
        self.assertTrue(project_document_mentions(text_only, "secondary_cabinet"))

        zero_qty_table_row = ProjectDocumentContext(
            project_root="D:/project",
            expected_counts={"secondary_cabinet": 0.0},
            text_snippets=("二次柜 | 0 | 台",),
        )
        self.assertFalse(project_document_mentions(zero_qty_table_row, "secondary_cabinet"))

        zero_qty_with_model_code = ProjectDocumentContext(
            project_root="D:/project",
            expected_counts={"secondary_cabinet": 0.0},
            text_snippets=("二次柜 | DK-1 | 0 | 台",),
        )
        self.assertFalse(project_document_mentions(zero_qty_with_model_code, "secondary_cabinet"))


def _write_docx(path: Path, *, rows: tuple[tuple[str, ...], ...], extra_paragraphs: tuple[str, ...] = ()) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("1.2 建设规模")
    for paragraph in extra_paragraphs:
        doc.add_paragraph(paragraph)
    table = doc.add_table(rows=len(rows) + 1, cols=6)
    headers = ("项目", "项目", "项目", "可研批复规模", "设计建设规模", "对比分析(±%)")
    for index, value in enumerate(headers):
        table.cell(0, index).text = value
    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    doc.save(path)


def _write_xlsx(path: Path, *, rows: tuple[tuple[str, ...], ...]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(list(row))
    workbook.save(path)


_DXF_TEXT_ONLY = """0
SECTION
2
ENTITIES
0
TEXT
10
0
20
0
1
智能网关
0
ENDSEC
0
EOF
"""
