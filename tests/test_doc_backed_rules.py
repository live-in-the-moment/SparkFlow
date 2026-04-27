from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from openpyxl import Workbook

from backend.core import audit_file
from backend.model.types import Device, Point2D, ProjectDocumentContext, SystemModel
from backend.rules.project_rules import _count_drawing_occurrences


class DocumentBackedRuleTests(unittest.TestCase):
    def test_stategrid_strict_v2_emits_document_backed_issues(self) -> None:
        with tempfile.TemporaryDirectory() as td:
            project_root = Path(td) / "工程A"
            drawing_path = project_root / "附件3 施工图" / "一次系统图.dxf"
            out_dir = Path(td) / "out"
            drawing_path.parent.mkdir(parents=True, exist_ok=True)
            drawing_path.write_text(_DXF_WITH_COUNTS, encoding="utf-8")

            _write_docx(
                project_root / "附件1 施工图设计说明书_工程A.docx",
                rows=(
                    ("低压配电网", "配变数量（台）", "配变数量（台）", "2", "2", "0"),
                    ("智能化", "配电智能网关", "配电智能网关", "2", "2", "0"),
                    ("配网自动化", "DTU（台）", "DTU（台）", "1", "1", "0"),
                ),
            )
            _write_xlsx(
                project_root / "附件4 主要设备材料清册.xlsx",
                rows=(
                    ("设备名称", "数量", "单位"),
                    ("综合测控通信单元", "1", "台"),
                    ("二次柜", "2", "台"),
                ),
            )
            tower_xls = project_root / "附件3 施工图" / "10kV及以下杆（塔）明细表-工程A.xls"
            tower_xls.write_bytes(b"stub")

            repo_ruleset = Path(__file__).resolve().parents[1] / "rulesets" / "stategrid_peidian_strict"
            with patch(
                "backend.project_docs._read_xls_rows_via_excel",
                return_value=[
                    ["名称", "数量", "单位"],
                    ["电杆", "3", "根"],
                ],
            ):
                result = audit_file(drawing_path, out_dir, ruleset_dir=repo_ruleset)

            report = json.loads(result.report_json_path.read_text(encoding="utf-8"))
            issue_ids = {issue["rule_id"] for issue in report["issues"]}

            self.assertEqual(report["rule_version"], "stategrid_peidian_strict_v2")
            self.assertIn("project.smart_gateway.count_mismatch", issue_ids)
            self.assertIn("project.distribution_transformer.count_mismatch", issue_ids)
            self.assertIn("project.pole_tower.count_mismatch", issue_ids)
            self.assertIn("project.dtu.count_mismatch", issue_ids)
            self.assertIn("project.measurement_comm_unit.missing_presence", issue_ids)
            self.assertIn("project.secondary_cabinet.count_mismatch", issue_ids)
            self.assertIn("project.secondary_cabinet.missing_quantity_note", issue_ids)
            self.assertFalse(report["passed"])
            self.assertTrue(report["summary"]["project_documents"]["enabled"])

    def test_distribution_transformer_count_excludes_ct_transformers(self) -> None:
        model = SystemModel(
            devices=(
                Device(
                    id="dev-ct",
                    position=Point2D(0.0, 0.0),
                    label="TA1",
                    device_type="transformer",
                    source_entity_ids=("text:ct",),
                ),
            ),
            texts=(
                ("text:ct", Point2D(0.0, 0.0), "电流互感器"),
            ),
            project_documents=ProjectDocumentContext(
                project_root="D:/project",
                expected_counts={"distribution_transformer": 1.0},
            ),
        )

        self.assertEqual(_count_drawing_occurrences(model, "distribution_transformer"), 0)

    def test_distribution_transformer_count_avoids_double_counting_true_overlap(self) -> None:
        model = SystemModel(
            devices=(
                Device(
                    id="dev-transformer",
                    position=Point2D(0.0, 0.0),
                    label="1#公用台变",
                    device_type="transformer",
                    source_entity_ids=("insert:1", "text:merged"),
                ),
            ),
            texts=(
                ("text:merged", Point2D(0.0, 0.0), "公用台变"),
                ("text:separate", Point2D(20.0, 0.0), "黄沙公用台变"),
            ),
            project_documents=ProjectDocumentContext(
                project_root="D:/project",
                expected_counts={"distribution_transformer": 2.0},
            ),
        )

        self.assertEqual(_count_drawing_occurrences(model, "distribution_transformer"), 2)

    def test_distribution_transformer_count_does_not_double_count_insert_backed_label_text(self) -> None:
        model = SystemModel(
            devices=(
                Device(
                    id="dev-insert-transformer",
                    position=Point2D(0.0, 0.0),
                    label="黄沙公用台变",
                    block_name="TR-01",
                    device_type="transformer",
                    source_entity_ids=("insert:1",),
                ),
            ),
            texts=(
                ("text:label", Point2D(1.0, 0.0), "黄沙公用台变"),
            ),
            project_documents=ProjectDocumentContext(
                project_root="D:/project",
                expected_counts={"distribution_transformer": 1.0},
            ),
        )

        self.assertEqual(_count_drawing_occurrences(model, "distribution_transformer"), 1)


def _write_docx(path: Path, *, rows: tuple[tuple[str, ...], ...]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("1.2 建设规模")
    table = doc.add_table(rows=len(rows) + 1, cols=6)
    headers = ("项目", "项目", "项目", "可研批复规模", "设计建设规模", "对比分析(±%)")
    for index, value in enumerate(headers):
        table.cell(0, index).text = value
    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    doc.save(path)


def _write_xlsx(path: Path, *, rows: tuple[tuple[str, ...], ...]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(list(row))
    workbook.save(path)


_DXF_WITH_COUNTS = """0
SECTION
2
ENTITIES
0
TEXT
10
0
20
0
1
智能网关
0
TEXT
10
10
20
0
1
黄沙公用台变
0
TEXT
10
20
20
0
1
N1杆
0
TEXT
10
30
20
0
1
N2杆
0
TEXT
10
40
20
0
1
二次柜
0
ENDSEC
0
EOF
"""
