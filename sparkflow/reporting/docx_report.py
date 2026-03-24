from __future__ import annotations

import json
from collections.abc import Iterable
from pathlib import Path

from ..contracts import AuditReport, Issue


def write_docx_report(report: AuditReport, out_path: Path) -> Path:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_paragraph("SparkFlow 审图报告")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"created_at: {report.created_at}")
    doc.add_paragraph(f"input_path: {report.input_path}")
    doc.add_paragraph(f"input_sha256: {report.input_sha256}")
    doc.add_paragraph(f"parser: {report.parser}")
    doc.add_paragraph(f"rule_version: {report.rule_version}")
    doc.add_paragraph(f"passed: {str(report.passed).lower()}")

    if report.summary:
        doc.add_heading("Summary", level=1)
        _kv_table(doc, [(str(k), _short_json(v)) for k, v in report.summary.items()])

    if report.artifacts:
        doc.add_heading("Artifacts", level=1)
        _kv_table(doc, [(str(k), str(v)) for k, v in report.artifacts.items()])

    doc.add_heading("Issues", level=1)
    if not report.issues:
        doc.add_paragraph("无问题。")
    else:
        _issue_table(doc, report.issues)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _kv_table(doc, rows: list[tuple[str, str]]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "字段"
    hdr[1].text = "值"
    for k, v in rows:
        r = table.add_row().cells
        r[0].text = k
        r[1].text = v


def _issue_table(doc, issues: Iterable[Issue]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "severity"
    hdr[2].text = "rule_id"
    hdr[3].text = "message"
    hdr[4].text = "refs"
    hdr[5].text = "refs(x,y)"
    for i, it in enumerate(list(issues), start=1):
        r = table.add_row().cells
        r[0].text = str(i)
        r[1].text = str(it.severity.value)
        r[2].text = str(it.rule_id)
        r[3].text = str(it.message)
        ref_parts = []
        xy_parts = []
        for ref in it.refs[:3]:
            ref_parts.append(f"{ref.kind}:{ref.id}")
            extra = ref.extra or {}
            x = extra.get("x")
            y = extra.get("y")
            if x is not None and y is not None:
                xy_parts.append(f"({x},{y})")
        r[4].text = " ".join(ref_parts)
        r[5].text = " ".join(xy_parts)


def _short_json(v: object) -> str:
    try:
        s = json.dumps(v, ensure_ascii=False)
    except Exception:
        s = str(v)
    if len(s) > 1600:
        return s[:1600] + "..."
    return s

