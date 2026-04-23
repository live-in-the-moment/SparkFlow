from __future__ import annotations

import os
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path, PurePosixPath
from typing import Iterable
import xml.etree.ElementTree as ET

from docx import Document

from .model.types import ProjectDocumentContext, ProjectDocumentFact, ProjectDocumentSource

_XLSX_MAIN_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
_XLSX_DOCUMENT_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_XLSX_PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_XLSX_NS = {
    "main": _XLSX_MAIN_NS,
    "rel": _XLSX_DOCUMENT_REL_NS,
    "pkg": _XLSX_PACKAGE_REL_NS,
}

_NOTE_TOKENS = ("工程量", "工程量说明", "设备选型", "材料清册")
_NOISE_TOKENS = (
    "工程名称",
    "项目编码",
    "容量",
    "状态量",
    "传感器",
    "配电箱",
    "标识",
    "标志",
    "材料",
    "安装图",
    "基础",
    "横担",
    "抱箍",
    "监测",
    "温控仪",
    "说明",
    "作业机械",
)


@dataclass(frozen=True)
class _ObjectSpec:
    key: str
    display_name: str
    aliases: tuple[str, ...]
    drawing_aliases: tuple[str, ...]
    docx_scale_labels: tuple[str, ...] = ()
    xlsx_exact_labels: tuple[str, ...] = ()


_OBJECT_SPECS = (
    _ObjectSpec(
        key="smart_gateway",
        display_name="智能网关",
        aliases=("配电智能网关", "智能网关"),
        drawing_aliases=("智能网关", "配电智能网关"),
        docx_scale_labels=("配电智能网关",),
        xlsx_exact_labels=("配电智能网关",),
    ),
    _ObjectSpec(
        key="distribution_transformer",
        display_name="台变",
        aliases=("配变数量", "配变", "台变", "公用台变", "变压器"),
        drawing_aliases=("台变", "公用台变", "配变", "变压器"),
        docx_scale_labels=("配变数量（台）",),
        xlsx_exact_labels=("台变", "公用台变"),
    ),
    _ObjectSpec(
        key="pole_tower",
        display_name="杆塔",
        aliases=("杆（塔）", "杆塔", "电杆"),
        drawing_aliases=("杆塔", "电杆", "杆"),
    ),
    _ObjectSpec(
        key="dtu",
        display_name="DTU",
        aliases=("DTU",),
        drawing_aliases=("DTU",),
        docx_scale_labels=("DTU（台）",),
        xlsx_exact_labels=("DTU",),
    ),
    _ObjectSpec(
        key="measurement_comm_unit",
        display_name="综合测控通信单元",
        aliases=("综合测控通信单元",),
        drawing_aliases=("综合测控通信单元", "综合测控"),
        xlsx_exact_labels=("综合测控通信单元",),
    ),
    _ObjectSpec(
        key="secondary_cabinet",
        display_name="二次柜",
        aliases=("二次柜",),
        drawing_aliases=("二次柜",),
        xlsx_exact_labels=("二次柜",),
    ),
)

_SPEC_BY_KEY = {spec.key: spec for spec in _OBJECT_SPECS}
_SOURCE_PATTERNS = (
    ("docx", "*设计说明书*.docx"),
    ("xlsx", "*主要设备材料清册*.xlsx"),
    ("xlsx", "*材料清册*.xlsx"),
    ("xls", "*杆（塔）明细表*.xls"),
    ("xlsx", "*杆（塔）明细表*.xlsx"),
)
_SOURCE_PRIORITY = {"docx": 0, "xlsx": 1, "xls": 1}


def build_project_document_context(input_path: Path) -> ProjectDocumentContext | None:
    project_root = _find_project_root(input_path.resolve())
    if project_root is None:
        return None

    sources = _discover_sources(project_root)
    if not sources:
        return None

    facts: list[ProjectDocumentFact] = []
    snippets: list[str] = []
    errors: list[str] = []
    for source in sources:
        try:
            source_facts, source_snippets = _extract_source(source)
            facts.extend(source_facts)
            snippets.extend(source_snippets)
        except Exception as exc:  # pragma: no cover - defensive
            errors.append(f"{source.kind}:{source.path.name}:{exc}")

    return ProjectDocumentContext(
        project_root=str(project_root),
        sources=tuple(ProjectDocumentSource(kind=source.kind, path=str(source.path)) for source in sources),
        facts=tuple(facts),
        expected_counts=_collapse_facts(facts),
        text_snippets=tuple(dict.fromkeys(snippets)),
        errors=tuple(errors),
    )


def project_document_aliases(key: str) -> tuple[str, ...]:
    spec = _SPEC_BY_KEY.get(key)
    return spec.drawing_aliases if spec is not None else ()


def project_document_note_exists(context: ProjectDocumentContext | None, key: str) -> bool:
    if context is None:
        return False
    spec = _SPEC_BY_KEY.get(key)
    if spec is None:
        return False
    for snippet in context.text_snippets:
        normalized = _normalize_text(snippet)
        if not any(alias in normalized for alias in spec.aliases):
            continue
        if any(token in normalized for token in _NOTE_TOKENS):
            return True
    return False


def project_document_mentions(context: ProjectDocumentContext | None, key: str) -> bool:
    if context is None:
        return False
    if context.expected_counts.get(key, 0) > 0:
        return True
    spec = _SPEC_BY_KEY.get(key)
    if spec is None:
        return False
    return any(_snippet_mentions_object(snippet, spec.aliases) for snippet in context.text_snippets)


def project_document_display_name(key: str) -> str:
    spec = _SPEC_BY_KEY.get(key)
    return spec.display_name if spec is not None else key


def _find_project_root(input_path: Path) -> Path | None:
    candidates = [input_path.parent, *input_path.parents[:4]]
    best_candidate: Path | None = None
    best_score = 0
    for candidate in candidates:
        if not candidate.exists() or not candidate.is_dir():
            continue
        score = _source_score(candidate)
        if score > best_score:
            best_candidate = candidate
            best_score = score
    return best_candidate


def _source_score(candidate: Path) -> int:
    attachment_dirs = [child for child in candidate.iterdir() if child.is_dir() and child.name.startswith("附件")]
    score = 0
    for _, pattern in _SOURCE_PATTERNS:
        if next(candidate.glob(pattern), None) is not None:
            score += 1
            continue
        if any(next(child.rglob(pattern), None) is not None for child in attachment_dirs):
            score += 1
    return score


@dataclass(frozen=True)
class _SourceFile:
    kind: str
    path: Path


def _discover_sources(project_root: Path) -> list[_SourceFile]:
    seen: set[Path] = set()
    sources: list[_SourceFile] = []
    for kind, pattern in _SOURCE_PATTERNS:
        for path in sorted(project_root.rglob(pattern)):
            if not path.is_file():
                continue
            if path in seen:
                continue
            seen.add(path)
            sources.append(_SourceFile(kind=kind, path=path.resolve()))
    return sources


def _extract_source(source: _SourceFile) -> tuple[list[ProjectDocumentFact], list[str]]:
    if source.kind == "docx":
        return _extract_docx(source.path)
    if source.kind == "xlsx":
        return _extract_xlsx(source.path)
    if source.kind == "xls":
        return _extract_xls(source.path)
    return [], []


def _extract_docx(path: Path) -> tuple[list[ProjectDocumentFact], list[str]]:
    doc = Document(path)
    snippets = [text for text in (_normalize_space(p.text) for p in doc.paragraphs) if text]
    facts: list[ProjectDocumentFact] = []
    for table in doc.tables:
        rows = [[_normalize_space(cell.text) for cell in row.cells] for row in table.rows]
        facts.extend(_extract_docx_scale_facts(rows, source_path=path))
        for row in rows:
            joined = " | ".join(cell for cell in row if cell)
            if joined:
                snippets.append(joined)
    return facts, snippets


def _extract_xlsx(path: Path) -> tuple[list[ProjectDocumentFact], list[str]]:
    facts: list[ProjectDocumentFact] = []
    snippets: list[str] = []
    if "杆（塔）明细表" in path.name or "杆塔明细表" in path.name:
        return _extract_pole_table_xlsx(path)
    for _, rows in _read_xlsx_sheets(path):
        facts.extend(_extract_xlsx_exact_facts(rows, source_path=path))
        for row in rows:
            joined = " | ".join(cell for cell in row if cell)
            if joined:
                snippets.append(joined)
    return facts, snippets


def _extract_xls(path: Path) -> tuple[list[ProjectDocumentFact], list[str]]:
    rows = _read_xls_rows(path)
    if "杆（塔）明细表" in path.name or "杆塔明细表" in path.name:
        facts = _extract_pole_table_rows(rows, source_kind="xls", source_path=path)
    else:
        facts = _extract_xlsx_exact_facts(rows, source_path=path, source_kind="xls")
    snippets = [" | ".join(cell for cell in row if cell) for row in rows if any(cell for cell in row)]
    return facts, snippets


def _snippet_mentions_object(snippet: str, aliases: tuple[str, ...]) -> bool:
    normalized = _normalize_text(snippet)
    if not any(alias in normalized for alias in aliases):
        return False
    if "|" not in snippet:
        return True
    numeric_values = [_parse_pure_numeric_cell(cell) for cell in snippet.split("|")]
    numbers = [value for value in numeric_values if value is not None]
    if numbers and all(value == 0 for value in numbers):
        return False
    return True


def _parse_pure_numeric_cell(cell: str) -> float | None:
    normalized = _normalize_space(cell).replace(",", "")
    if not normalized:
        return None
    if not re.fullmatch(r"-?\d+(?:\.\d+)?", normalized):
        return None
    try:
        return float(normalized)
    except ValueError:
        return None


def _extract_docx_scale_facts(rows: Iterable[Iterable[str]], *, source_path: Path) -> list[ProjectDocumentFact]:
    normalized_rows = [tuple(_normalize_space(cell) for cell in row) for row in rows]
    facts: list[ProjectDocumentFact] = []
    design_idx = None
    for row in normalized_rows:
        if not any(row):
            continue
        if design_idx is None:
            design_idx = _find_design_scale_column(row)
        if design_idx is None:
            continue
        raw_label = row[1] if len(row) > 1 else row[0]
        normalized_label = _normalize_text(raw_label)
        for spec in _OBJECT_SPECS:
            if not spec.docx_scale_labels:
                continue
            if not any(normalized_label == _normalize_text(label) for label in spec.docx_scale_labels):
                continue
            value = _extract_row_quantity(row, quantity_idx=None, design_idx=design_idx)
            if value is None:
                continue
            facts.append(
                ProjectDocumentFact(
                    key=spec.key,
                    display_name=spec.display_name,
                    value=value,
                    source_kind="docx",
                    source_path=str(source_path),
                    raw_label=raw_label,
                    unit=_infer_unit(row),
                )
            )
    return facts


def _extract_xlsx_exact_facts(
    rows: Iterable[Iterable[str]],
    *,
    source_path: Path,
    source_kind: str = "xlsx",
) -> list[ProjectDocumentFact]:
    normalized_rows = [tuple(_normalize_space(cell) for cell in row) for row in rows]
    facts: list[ProjectDocumentFact] = []
    quantity_idx = None
    for row in normalized_rows:
        if not any(row):
            continue
        if quantity_idx is None:
            quantity_idx = _find_quantity_column(row)
        if quantity_idx is None:
            continue
        for cell in row:
            normalized = _normalize_text(cell)
            if not normalized:
                continue
            for spec in _OBJECT_SPECS:
                if not spec.xlsx_exact_labels:
                    continue
                if not any(normalized == _normalize_text(label) for label in spec.xlsx_exact_labels):
                    continue
                value = _extract_row_quantity(row, quantity_idx=quantity_idx, design_idx=None)
                if value is None:
                    continue
                facts.append(
                    ProjectDocumentFact(
                        key=spec.key,
                        display_name=spec.display_name,
                        value=value,
                        source_kind=source_kind,
                        source_path=str(source_path),
                        raw_label=cell,
                        unit=_infer_unit(row),
                    )
                )
    return facts


def _extract_pole_table_xlsx(path: Path) -> tuple[list[ProjectDocumentFact], list[str]]:
    snippets: list[str] = []
    rows: list[list[str]] = []
    for _, sheet_rows in _read_xlsx_sheets(path):
        rows.extend(sheet_rows)
        for row in sheet_rows:
            joined = " | ".join(cell for cell in row if cell)
            if joined:
                snippets.append(joined)
    return _extract_pole_table_rows(rows, source_kind="xlsx", source_path=path), snippets


def _extract_pole_table_rows(
    rows: Iterable[Iterable[str]],
    *,
    source_kind: str,
    source_path: Path,
) -> list[ProjectDocumentFact]:
    normalized_rows = [tuple(_normalize_space(cell) for cell in row) for row in rows]
    pole_ids: list[str] = []
    for cells in normalized_rows:
        if not cells:
            continue
        first = cells[0] if cells else ""
        if re.fullmatch(r"[A-Z]+\d+", first):
            pole_ids.append(first)
    if pole_ids:
        return [
            ProjectDocumentFact(
                key="pole_tower",
                display_name=project_document_display_name("pole_tower"),
                value=float(len(set(pole_ids))),
                source_kind=source_kind,
                source_path=str(source_path),
                raw_label="杆（塔）明细表",
                unit="基",
            )
        ]

    quantity_idx = None
    for row in normalized_rows:
        if quantity_idx is None:
            quantity_idx = _find_quantity_column(row)
        if quantity_idx is None:
            continue
        for cell in row:
            if _normalize_text(cell) not in {_normalize_text("电杆"), _normalize_text("杆塔"), _normalize_text("杆（塔）")}:
                continue
            value = _extract_row_quantity(row, quantity_idx=quantity_idx, design_idx=None)
            if value is None:
                continue
            return [
                ProjectDocumentFact(
                    key="pole_tower",
                    display_name=project_document_display_name("pole_tower"),
                    value=value,
                    source_kind=source_kind,
                    source_path=str(source_path),
                    raw_label=cell,
                    unit=_infer_unit(row),
                )
            ]
    return []


def _find_quantity_column(row: Iterable[str]) -> int | None:
    for idx, cell in enumerate(row):
        if "数量" in cell:
            return idx
    return None


def _find_design_scale_column(row: Iterable[str]) -> int | None:
    for idx, cell in enumerate(row):
        if "设计建设规模" in cell:
            return idx
    return None


def _match_row_object(row: Iterable[str]) -> tuple[_ObjectSpec, str] | None:
    for cell in row:
        normalized = _normalize_text(cell)
        if not normalized:
            continue
        for spec in _OBJECT_SPECS:
            exact_alias = any(normalized == _normalize_text(alias) for alias in spec.aliases)
            if not exact_alias:
                if len(normalized) > 24:
                    continue
                if any(token in normalized for token in _NOISE_TOKENS):
                    continue
            for alias in spec.aliases:
                if alias in normalized:
                    return spec, cell
    return None


def _extract_row_quantity(
    row: Iterable[str],
    *,
    quantity_idx: int | None,
    design_idx: int | None,
) -> float | None:
    cells = list(row)
    preferred_indexes = []
    if design_idx is not None and design_idx < len(cells):
        preferred_indexes.append(design_idx)
    if quantity_idx is not None and quantity_idx < len(cells):
        preferred_indexes.append(quantity_idx)
    preferred_indexes.extend(idx for idx in range(len(cells)) if idx not in preferred_indexes)
    for idx in preferred_indexes:
        value = _parse_number(cells[idx])
        if value is not None:
            return value
    return None


def _infer_unit(row: Iterable[str]) -> str | None:
    for cell in row:
        normalized = _normalize_text(cell)
        if any(token in normalized for token in ("（台）", "(台)", "（个）", "(个)", "（根）", "(根)", "台", "个", "根")):
            return cell
    return None


def _collapse_facts(facts: Iterable[ProjectDocumentFact]) -> dict[str, float]:
    best: dict[str, ProjectDocumentFact] = {}
    for fact in facts:
        current = best.get(fact.key)
        if current is None:
            best[fact.key] = fact
            continue
        current_priority = _fact_priority(current)
        fact_priority = _fact_priority(fact)
        if fact_priority < current_priority:
            best[fact.key] = fact
            continue
        if fact_priority == current_priority and fact.value > 0 and current.value <= 0:
            best[fact.key] = fact
    return {key: fact.value for key, fact in best.items()}


def _fact_priority(fact: ProjectDocumentFact) -> tuple[int, int]:
    source_priority = _SOURCE_PRIORITY.get(fact.source_kind, 99)
    if fact.key == "pole_tower" and fact.source_kind in {"xls", "xlsx"} and "杆（塔）明细表" in fact.raw_label:
        return (0, 0)
    if fact.key in {"smart_gateway", "distribution_transformer", "dtu"} and fact.source_kind == "docx":
        return (0, 0)
    if fact.key in {"measurement_comm_unit", "secondary_cabinet"} and fact.source_kind == "xlsx":
        return (0, 0)
    return (source_priority + 5, 0)


def _parse_number(text: str) -> float | None:
    normalized = _normalize_space(text).replace(",", "")
    if not normalized:
        return None
    if normalized.endswith("%"):
        return None
    match = re.search(r"-?\d+(?:\.\d+)?", normalized)
    if match is None:
        return None
    try:
        return float(match.group(0))
    except ValueError:
        return None


def _normalize_space(text: object) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _normalize_text(text: object) -> str:
    return _normalize_space(text).replace(" ", "")


def _read_xls_rows(path: Path) -> list[list[str]]:
    try:
        return _read_xls_rows_via_excel(path)
    except Exception as exc:  # pragma: no cover - machine-specific fallback
        raise RuntimeError(f"读取 xls 失败：{exc}") from exc


def _read_xls_rows_via_excel(path: Path) -> list[list[str]]:
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    excel = None
    workbook = None
    try:
        excel = win32com.client.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        workbook = excel.Workbooks.Open(os.fspath(path))
        sheet = workbook.Worksheets(1)
        used = sheet.UsedRange.Value
        if used is None:
            return []
        if not isinstance(used, tuple):
            used = ((used,),)
        rows: list[list[str]] = []
        for row in used:
            if not isinstance(row, tuple):
                row = (row,)
            rows.append([_normalize_space(cell) for cell in row])
        return rows
    finally:
        if workbook is not None:
            workbook.Close(False)
        if excel is not None:
            excel.Quit()
        pythoncom.CoUninitialize()


def _read_xlsx_sheets(path: Path) -> list[tuple[str, list[list[str]]]]:
    try:
        with zipfile.ZipFile(path) as archive:
            sheet_paths = _resolve_xlsx_sheet_paths(archive)
            shared_strings = _load_xlsx_shared_strings(archive)
            sheets: list[tuple[str, list[list[str]]]] = []
            for sheet_name, sheet_path in sheet_paths:
                sheet_root = ET.fromstring(archive.read(sheet_path))
                rows: list[list[str]] = []
                for row in sheet_root.findall("./main:sheetData/main:row", _XLSX_NS):
                    row_values: dict[int, str] = {}
                    for cell in row.findall("main:c", _XLSX_NS):
                        column_index = _xlsx_column_index(cell.get("r"))
                        if column_index <= 0:
                            continue
                        row_values[column_index] = _xlsx_cell_text(cell, shared_strings)
                    max_column = max(row_values, default=0)
                    if max_column:
                        rows.append([_normalize_space(row_values.get(index, "")) for index in range(1, max_column + 1)])
                    else:
                        rows.append([])
                sheets.append((sheet_name, rows))
            return sheets
    except (KeyError, ET.ParseError, zipfile.BadZipFile, ValueError) as exc:
        raise ValueError(f"{path.name} 不是有效的 XLSX 文件。") from exc


def _resolve_xlsx_sheet_paths(archive: zipfile.ZipFile) -> list[tuple[str, str]]:
    workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
    rels_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
    rel_map = {
        rel.get("Id"): rel.get("Target", "")
        for rel in rels_root.findall("./pkg:Relationship", _XLSX_NS)
        if rel.get("Id")
    }
    paths: list[tuple[str, str]] = []
    for sheet in workbook_root.findall("./main:sheets/main:sheet", _XLSX_NS):
        name = sheet.get("name") or ""
        rel_id = sheet.get(f"{{{_XLSX_DOCUMENT_REL_NS}}}id")
        if not rel_id:
            continue
        target = rel_map.get(rel_id)
        if not target:
            continue
        paths.append((name, _normalize_zip_path("xl", target)))
    if not paths:
        raise ValueError("workbook 缺少工作表。")
    return paths


def _load_xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in archive.namelist():
        return []
    shared_root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    values: list[str] = []
    for item in shared_root.findall("./main:si", _XLSX_NS):
        values.append("".join(node.text or "" for node in item.findall(".//main:t", _XLSX_NS)))
    return values


def _normalize_zip_path(base_dir: str, target: str) -> str:
    if target.startswith("/"):
        target_path = PurePosixPath(target.lstrip("/"))
    else:
        target_path = PurePosixPath(base_dir) / PurePosixPath(target)
    parts: list[str] = []
    for part in target_path.parts:
        if part in {"", "."}:
            continue
        if part == "..":
            if parts:
                parts.pop()
            continue
        parts.append(part)
    return "/".join(parts)


def _xlsx_column_index(ref: str | None) -> int:
    if not ref:
        return 0
    letters = ""
    for ch in ref:
        if ch.isalpha():
            letters += ch
        else:
            break
    out = 0
    for ch in letters:
        out = out * 26 + (ord(ch.upper()) - 64)
    return out


def _xlsx_cell_text(cell: ET.Element, shared_strings: list[str]) -> str:
    cell_type = cell.get("t")
    if cell_type == "inlineStr":
        return "".join(node.text or "" for node in cell.findall(".//main:t", _XLSX_NS))
    value_node = cell.find("main:v", _XLSX_NS)
    if value_node is None or value_node.text is None:
        return ""
    text = value_node.text
    if cell_type == "s":
        try:
            return shared_strings[int(text)]
        except (IndexError, ValueError):
            return text
    return text
