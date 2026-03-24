from __future__ import annotations

import argparse
import json
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(prog="generate_audit_docx")
    ap.add_argument("--dwg", type=Path, required=True)
    ap.add_argument("--dxf", type=Path, required=True)
    ap.add_argument("--info-json", type=Path, required=True)
    ap.add_argument("--summary-json", type=Path, required=True)
    ap.add_argument("--audit-json", type=Path, required=True)
    ap.add_argument("--audit-md", type=Path, required=True)
    ap.add_argument("--out", type=Path, required=True)
    ap.add_argument("--max-issues", type=int, default=80)
    args = ap.parse_args(argv)

    info = json.loads(args.info_json.read_text(encoding="utf-8"))
    summary = json.loads(args.summary_json.read_text(encoding="utf-8"))
    audit = json.loads(args.audit_json.read_text(encoding="utf-8"))

    doc = Document()

    title = doc.add_paragraph("SparkFlow 单图审图解析报告")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("生成时间：").bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    _add_h1(doc, "一、输入与产物")
    _kv_table(
        doc,
        [
            ("DWG 输入", str(args.dwg.resolve())),
            ("DXF（转换产物）", str(args.dxf.resolve())),
            ("全量解析 JSON", str(args.info_json.resolve())),
            ("汇总解析 JSON", str(args.summary_json.resolve())),
            ("审图报告 JSON", str(args.audit_json.resolve())),
            ("审图报告 MD", str(args.audit_md.resolve())),
        ],
    )

    _add_h1(doc, "二、图纸全量信息摘要")
    _kv_table(
        doc,
        [
            ("parser_id", str(info.get("parser_id"))),
            ("entity_count", str(info.get("entity_count"))),
            ("bbox", json.dumps(info.get("bbox"), ensure_ascii=False)),
        ],
    )

    _add_h2(doc, "2.1 图元类型统计")
    _counter_table(doc, info.get("kinds") or {})

    _add_h2(doc, "2.2 图层统计")
    _counter_table(doc, info.get("layers") or {}, top_n=30)

    _add_h2(doc, "2.3 块名（INSERT）")
    _list(doc, summary.get("blocks") or [])

    _add_h2(doc, "2.4 图号/模块代号")
    _list(doc, summary.get("drawing_codes") or [])

    _add_h2(doc, "2.5 材料/强度/尺寸要点（用于审图判断）")
    _kv_table(
        doc,
        [
            ("混凝土等级", "、".join(summary.get("concretes") or []) or "（未识别）"),
            ("材料牌号", "、".join(summary.get("materials") or []) or "（未识别）"),
            ("尺寸/指标", "、".join(summary.get("dimensions") or []) or "（未识别）"),
        ],
    )

    _add_h1(doc, "三、审图结论与问题清单")
    passed = bool(audit.get("passed"))
    p = doc.add_paragraph()
    p.add_run("结论：").bold = True
    p.add_run("通过" if passed else "未通过")

    issues = list(audit.get("issues") or [])
    rule_counts = Counter(i.get("rule_id") for i in issues if i.get("rule_id"))
    severity_counts = Counter(i.get("severity") for i in issues if i.get("severity"))

    _add_h2(doc, "3.1 问题统计")
    _kv_table(
        doc,
        [
            ("issues_total", str(len(issues))),
            ("by_severity", json.dumps(dict(severity_counts), ensure_ascii=False)),
            ("by_rule", json.dumps(dict(rule_counts), ensure_ascii=False)),
        ],
    )

    _add_h2(doc, "3.2 问题明细（截取前 N 条）")
    max_issues = max(0, int(args.max_issues))
    _issue_table(doc, issues[:max_issues])

    doc.add_page_break()
    _add_h1(doc, "附录：文本全集（截取）")
    unique_texts = list(info.get("unique_texts") or [])
    _list(doc, unique_texts[:200])

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    print(str(args.out.resolve()))
    return 0


def _add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles["Heading 1"]


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles["Heading 2"]


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "字段"
    hdr[1].text = "值"
    for k, v in rows:
        r = table.add_row().cells
        r[0].text = str(k)
        r[1].text = str(v)


def _counter_table(doc: Document, mapping: dict, *, top_n: int | None = None) -> None:
    items = list(mapping.items())
    items.sort(key=lambda kv: (-int(kv[1]), str(kv[0])))
    if top_n is not None:
        items = items[:top_n]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "名称"
    hdr[1].text = "数量"
    for k, v in items:
        r = table.add_row().cells
        r[0].text = str(k)
        r[1].text = str(v)


def _list(doc: Document, items: list[str]) -> None:
    if not items:
        doc.add_paragraph("（无）")
        return
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def _issue_table(doc: Document, issues: list[dict]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "severity"
    hdr[2].text = "rule_id"
    hdr[3].text = "message"
    hdr[4].text = "refs(x,y)"
    for i, it in enumerate(issues, start=1):
        r = table.add_row().cells
        r[0].text = str(i)
        r[1].text = str(it.get("severity") or "")
        r[2].text = str(it.get("rule_id") or "")
        r[3].text = str(it.get("message") or "")
        refs = it.get("refs") or []
        xy = []
        for ref in refs[:3]:
            extra = ref.get("extra") or {}
            x = extra.get("x")
            y = extra.get("y")
            if x is None or y is None:
                continue
            xy.append(f"({x:.3f},{y:.3f})" if isinstance(x, (int, float)) and isinstance(y, (int, float)) else f"({x},{y})")
        r[4].text = " ".join(xy)


if __name__ == "__main__":
    raise SystemExit(main())

