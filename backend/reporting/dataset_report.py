from __future__ import annotations

import hashlib
import json
from collections import Counter
from dataclasses import dataclass
from pathlib import Path, PureWindowsPath

from ..contracts import Severity
from ..rules.knowledgebase import load_ruleset_dir

_DRAWING_CLASS_LABELS = {
    'supported_electrical': '已审计电气图纸',
    'geometry_only': '布置/几何图纸',
    'unsupported': '暂不支持图纸',
}

_STATUS_LABELS = {
    'passed': '通过',
    'failed': '未通过',
    'skipped': '跳过',
    'unprocessed': '未处理',
}


@dataclass(frozen=True)
class DatasetReportArtifacts:
    markdown_path: Path
    docx_path: Path


@dataclass(frozen=True)
class DatasetDrawingEntry:
    rel_path: str
    drawing_class: str
    reason: str
    status: str
    drawing_type_label: str
    issue_count: int
    issue_counts_by_rule: tuple[tuple[str, int], ...]
    sample_points: tuple[str, ...]
    report_md_rel: str | None
    report_docx_rel: str | None


@dataclass(frozen=True)
class DatasetReportModel:
    title: str
    created_at: str
    dataset_dir: str
    dataset_label: str
    rule_version: str
    strict_note: str | None
    counts: dict[str, int]
    selection_counts: dict[str, int]
    issues_by_rule: tuple[tuple[str, int], ...]
    elapsed_sec: float
    avg_file_sec: float
    dataset_summary_rel: str
    dataset_selection_rel: str
    audited_entries: tuple[DatasetDrawingEntry, ...]
    skipped_geometry_entries: tuple[DatasetDrawingEntry, ...]
    skipped_other_entries: tuple[DatasetDrawingEntry, ...]


def write_dataset_audit_report(
    run_dir: Path,
    *,
    ruleset_dir: Path | None = None,
    out_md: Path | None = None,
    out_docx: Path | None = None,
    title: str | None = None,
    dataset_label: str | None = None,
) -> DatasetReportArtifacts:
    run_dir = run_dir.resolve()
    model = build_dataset_audit_report_model(
        run_dir,
        ruleset_dir=ruleset_dir,
        title=title,
        dataset_label=dataset_label,
    )

    markdown_path = (out_md or (run_dir / 'final_audit_report.md')).resolve()
    docx_path = (out_docx or (run_dir / 'final_audit_report.docx')).resolve()

    markdown_path.parent.mkdir(parents=True, exist_ok=True)
    markdown_path.write_text(render_dataset_audit_report_markdown(model), encoding='utf-8')
    write_dataset_audit_report_docx(model, docx_path)
    return DatasetReportArtifacts(markdown_path=markdown_path, docx_path=docx_path)


def build_dataset_audit_report_model(
    run_dir: Path,
    *,
    ruleset_dir: Path | None = None,
    title: str | None = None,
    dataset_label: str | None = None,
) -> DatasetReportModel:
    summary_path = run_dir / 'dataset_summary.json'
    selection_path = run_dir / 'dataset_selection.json'
    if not summary_path.exists():
        raise FileNotFoundError(str(summary_path))
    if not selection_path.exists():
        raise FileNotFoundError(str(selection_path))

    summary = _read_json(summary_path)
    selection_rows = _read_json(selection_path)
    if not isinstance(summary, dict):
        raise ValueError('dataset_summary.json 必须是对象。')
    if not isinstance(selection_rows, list):
        raise ValueError('dataset_selection.json 必须是数组。')

    resolved_dataset_label = dataset_label or _path_leaf(str(summary.get('dataset_dir') or run_dir.name))
    strict_note = _build_strict_note(ruleset_dir)
    resolved_title = title or (
        f'{resolved_dataset_label} 严格审图最终报告'
        if strict_note
        else f'{resolved_dataset_label} 数据集审图最终报告'
    )

    entries = tuple(_load_entry(run_dir, row) for row in selection_rows if isinstance(row, dict))
    audited_entries = tuple(entry for entry in entries if entry.drawing_class == 'supported_electrical')
    skipped_geometry_entries = tuple(entry for entry in entries if entry.drawing_class == 'geometry_only')
    skipped_other_entries = tuple(
        entry for entry in entries if entry.drawing_class not in {'supported_electrical', 'geometry_only'}
    )

    timing = summary.get('timing') or {}
    counts = summary.get('counts') or {}
    selection_counts = summary.get('selection_counts') or {}
    issues_by_rule = summary.get('issues_by_rule') or {}

    return DatasetReportModel(
        title=resolved_title,
        created_at=str(summary.get('created_at') or ''),
        dataset_dir=str(summary.get('dataset_dir') or ''),
        dataset_label=resolved_dataset_label,
        rule_version=str(summary.get('rule_version') or ''),
        strict_note=str(strict_note) if strict_note else None,
        counts={key: int(counts.get(key, 0) or 0) for key in ('passed', 'failed', 'skipped', 'unprocessed')},
        selection_counts={
            key: int(selection_counts.get(key, 0) or 0)
            for key in ('supported_electrical', 'geometry_only', 'unsupported')
        },
        issues_by_rule=tuple(
            sorted(
                ((str(rule_id), int(count)) for rule_id, count in issues_by_rule.items()),
                key=lambda item: (-item[1], item[0]),
            )
        ),
        elapsed_sec=float(timing.get('elapsed_sec', 0.0) or 0.0),
        avg_file_sec=float(timing.get('avg_file_sec', 0.0) or 0.0),
        dataset_summary_rel=summary_path.relative_to(run_dir).as_posix(),
        dataset_selection_rel=selection_path.relative_to(run_dir).as_posix(),
        audited_entries=audited_entries,
        skipped_geometry_entries=skipped_geometry_entries,
        skipped_other_entries=skipped_other_entries,
    )


def render_dataset_audit_report_markdown(model: DatasetReportModel) -> str:
    lines: list[str] = []
    lines.append(f'# {model.title}')
    lines.append('')
    lines.append('## 1. 运行概览')
    lines.append('')
    lines.append(f'- 数据集目录：{model.dataset_dir}')
    lines.append(f'- 生成时间：{model.created_at}')
    lines.append(f'- 规则版本：{model.rule_version}')
    if model.strict_note:
        lines.append(f'- 严格规则：{model.strict_note}')
    lines.append(f'- 数据集汇总：`{model.dataset_summary_rel}`')
    lines.append(f'- 选图清单：`{model.dataset_selection_rel}`')
    lines.append('')
    lines.append('## 2. 结果总览')
    lines.append('')
    lines.append(
        f"- 已审计电气图纸：{model.selection_counts.get('supported_electrical', 0)} "
        f"（通过 {model.counts.get('passed', 0)}，未通过 {model.counts.get('failed', 0)}，未处理 {model.counts.get('unprocessed', 0)}）"
    )
    lines.append(
        f"- 跳过图纸：{model.counts.get('skipped', 0)} "
        f"（布置/几何 {model.selection_counts.get('geometry_only', 0)}，暂不支持 {model.selection_counts.get('unsupported', 0)}）"
    )
    lines.append(f'- 总耗时：{model.elapsed_sec:.3f} 秒')
    lines.append(f'- 平均单图耗时：{model.avg_file_sec:.3f} 秒')
    lines.append('')
    lines.append('## 3. 严格运行问题汇总')
    lines.append('')
    if not model.issues_by_rule:
        lines.append('- 未发现问题。')
    else:
        for rule_id, count in model.issues_by_rule:
            lines.append(f'- {rule_id}: {count}')
    lines.append('')
    lines.append('## 4. 已审计电气图纸')
    lines.append('')
    if not model.audited_entries:
        lines.append('无。')
    else:
        for index, entry in enumerate(model.audited_entries, start=1):
            lines.append(f'{index}. {entry.rel_path}')
            lines.append(f'   - 图纸类别：{entry.drawing_type_label}')
            lines.append(f'   - 审图结果：{_STATUS_LABELS.get(entry.status, entry.status)}')
            lines.append(f'   - 问题数：{entry.issue_count}')
            lines.append(f'   - 规则命中：{_format_rule_hits(entry.issue_counts_by_rule)}')
            lines.append(f'   - 代表坐标：{_format_points(entry.sample_points)}')
            if entry.report_md_rel or entry.report_docx_rel:
                lines.append(f'   - 单图报告：{_format_report_refs(entry.report_md_rel, entry.report_docx_rel)}')
    lines.append('')
    lines.append('## 5. 跳过的布置/几何图纸')
    lines.append('')
    if not model.skipped_geometry_entries:
        lines.append('无。')
    else:
        for index, entry in enumerate(model.skipped_geometry_entries, start=1):
            lines.append(
                f"{index}. {entry.rel_path} | {_DRAWING_CLASS_LABELS.get(entry.drawing_class, entry.drawing_class)} "
                f"| {entry.reason}"
            )
            if entry.report_md_rel or entry.report_docx_rel:
                lines.append(f'   - 单图报告：{_format_report_refs(entry.report_md_rel, entry.report_docx_rel)}')
    lines.append('')
    lines.append('## 6. 其他跳过图纸')
    lines.append('')
    if not model.skipped_other_entries:
        lines.append('无。')
    else:
        for index, entry in enumerate(model.skipped_other_entries, start=1):
            lines.append(
                f"{index}. {entry.rel_path} | {_DRAWING_CLASS_LABELS.get(entry.drawing_class, entry.drawing_class)} "
                f"| {entry.reason}"
            )
            if entry.report_md_rel or entry.report_docx_rel:
                lines.append(f'   - 单图报告：{_format_report_refs(entry.report_md_rel, entry.report_docx_rel)}')
    lines.append('')
    lines.append('## 7. 结论')
    lines.append('')
    if model.counts.get('unprocessed', 0) == 0:
        lines.append('- 本次严格跑批已完成，未处理图纸数为 0。')
    else:
        lines.append(f"- 本次严格跑批存在 {model.counts.get('unprocessed', 0)} 份未处理图纸，需补充排查。")
    if model.counts.get('failed', 0) > 0:
        lines.append(f"- 严格规则下共有 {model.counts.get('failed', 0)} 份电气图纸未通过。")
    else:
        lines.append('- 严格规则下所有电气图纸均通过。')
    lines.append('')
    return '\n'.join(lines)


def write_dataset_audit_report_docx(model: DatasetReportModel, out_path: Path) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    _configure_docx_styles(doc, qn, Pt)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = doc.styles['Title']
    title.add_run(model.title)

    doc.add_heading('1. 运行概览', level=1)
    _doc_line(doc, f'数据集目录：{model.dataset_dir}')
    _doc_line(doc, f'生成时间：{model.created_at}')
    _doc_line(doc, f'规则版本：{model.rule_version}')
    if model.strict_note:
        _doc_line(doc, f'严格规则：{model.strict_note}')
    _doc_line(doc, f'数据集汇总：{model.dataset_summary_rel}')
    _doc_line(doc, f'选图清单：{model.dataset_selection_rel}')

    doc.add_heading('2. 结果总览', level=1)
    _doc_line(
        doc,
        f"已审计电气图纸：{model.selection_counts.get('supported_electrical', 0)} "
        f"（通过 {model.counts.get('passed', 0)}，未通过 {model.counts.get('failed', 0)}，未处理 {model.counts.get('unprocessed', 0)}）",
    )
    _doc_line(
        doc,
        f"跳过图纸：{model.counts.get('skipped', 0)} "
        f"（布置/几何 {model.selection_counts.get('geometry_only', 0)}，暂不支持 {model.selection_counts.get('unsupported', 0)}）",
    )
    _doc_line(doc, f'总耗时：{model.elapsed_sec:.3f} 秒')
    _doc_line(doc, f'平均单图耗时：{model.avg_file_sec:.3f} 秒')

    doc.add_heading('3. 严格运行问题汇总', level=1)
    if not model.issues_by_rule:
        _doc_line(doc, '未发现问题。')
    else:
        for rule_id, count in model.issues_by_rule:
            _doc_line(doc, f'{rule_id}: {count}')

    doc.add_heading('4. 已审计电气图纸', level=1)
    if not model.audited_entries:
        _doc_line(doc, '无。')
    else:
        for index, entry in enumerate(model.audited_entries, start=1):
            doc.add_heading(f'{index}. {entry.rel_path}', level=2)
            _doc_line(doc, f'图纸类别：{entry.drawing_type_label}')
            _doc_line(doc, f'审图结果：{_STATUS_LABELS.get(entry.status, entry.status)}')
            _doc_line(doc, f'问题数：{entry.issue_count}')
            _doc_line(doc, f'规则命中：{_format_rule_hits(entry.issue_counts_by_rule)}')
            _doc_line(doc, f'代表坐标：{_format_points(entry.sample_points)}')
            if entry.report_md_rel or entry.report_docx_rel:
                _doc_line(doc, f'单图报告：{_format_report_refs(entry.report_md_rel, entry.report_docx_rel)}')

    doc.add_heading('5. 跳过的布置/几何图纸', level=1)
    if not model.skipped_geometry_entries:
        _doc_line(doc, '无。')
    else:
        for index, entry in enumerate(model.skipped_geometry_entries, start=1):
            _doc_line(
                doc,
                f"{index}. {entry.rel_path} | {_DRAWING_CLASS_LABELS.get(entry.drawing_class, entry.drawing_class)} | {entry.reason}",
            )
            if entry.report_md_rel or entry.report_docx_rel:
                _doc_line(doc, f'   单图报告：{_format_report_refs(entry.report_md_rel, entry.report_docx_rel)}')

    doc.add_heading('6. 其他跳过图纸', level=1)
    if not model.skipped_other_entries:
        _doc_line(doc, '无。')
    else:
        for index, entry in enumerate(model.skipped_other_entries, start=1):
            _doc_line(
                doc,
                f"{index}. {entry.rel_path} | {_DRAWING_CLASS_LABELS.get(entry.drawing_class, entry.drawing_class)} | {entry.reason}",
            )
            if entry.report_md_rel or entry.report_docx_rel:
                _doc_line(doc, f'   单图报告：{_format_report_refs(entry.report_md_rel, entry.report_docx_rel)}')

    doc.add_heading('7. 结论', level=1)
    if model.counts.get('unprocessed', 0) == 0:
        _doc_line(doc, '本次严格跑批已完成，未处理图纸数为 0。')
    else:
        _doc_line(doc, f"本次严格跑批存在 {model.counts.get('unprocessed', 0)} 份未处理图纸，需补充排查。")
    if model.counts.get('failed', 0) > 0:
        _doc_line(doc, f"严格规则下共有 {model.counts.get('failed', 0)} 份电气图纸未通过。")
    else:
        _doc_line(doc, '严格规则下所有电气图纸均通过。')

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _configure_docx_styles(doc, qn, Pt) -> None:
    for style_name, font_name, size in (
        ('Normal', '宋体', 10.5),
        ('Title', '微软雅黑', 18),
        ('Heading 1', '微软雅黑', 14),
        ('Heading 2', '微软雅黑', 12),
    ):
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), font_name)


def _doc_line(doc, text: str) -> None:
    doc.add_paragraph(text)


def _load_entry(run_dir: Path, row: dict[str, object]) -> DatasetDrawingEntry:
    rel_path = str(row.get('rel_path') or '')
    file_dir = _dataset_file_out_dir(run_dir, rel_path)
    report_path = file_dir / 'report.json'
    report = _read_json(report_path) if report_path.exists() else {}
    summary = report.get('summary') if isinstance(report, dict) else {}
    classification = summary.get('classification') if isinstance(summary, dict) else {}
    issues = report.get('issues') if isinstance(report, dict) else []
    if not isinstance(issues, list):
        issues = []
    issue_counts = Counter(
        str(issue.get('rule_id'))
        for issue in issues
        if isinstance(issue, dict) and issue.get('rule_id')
    )
    report_md_path = file_dir / 'report.md'
    report_docx_path = file_dir / 'report.docx'
    return DatasetDrawingEntry(
        rel_path=rel_path,
        drawing_class=str(row.get('drawing_class') or classification.get('drawing_class') or 'unknown'),
        reason=str(row.get('reason') or classification.get('reason') or ''),
        status=str(row.get('status') or ''),
        drawing_type_label=str(classification.get('drawing_type_label') or '未标注'),
        issue_count=len(issues),
        issue_counts_by_rule=tuple(sorted(issue_counts.items(), key=lambda item: (-item[1], item[0]))),
        sample_points=_collect_sample_points(issues),
        report_md_rel=report_md_path.relative_to(run_dir).as_posix() if report_md_path.exists() else None,
        report_docx_rel=report_docx_path.relative_to(run_dir).as_posix() if report_docx_path.exists() else None,
    )


def _dataset_file_out_dir(run_dir: Path, rel_path: str) -> Path:
    path = Path(*rel_path.split('/'))
    token = hashlib.sha1(rel_path.encode('utf-8')).hexdigest()[:8]
    return run_dir / 'files' / path.parent / f'{path.stem}__{token}'


def _collect_sample_points(issues: list[dict[str, object]]) -> tuple[str, ...]:
    points: list[str] = []
    for issue in issues:
        if not isinstance(issue, dict):
            continue
        refs = issue.get('refs')
        if not isinstance(refs, list):
            continue
        for ref in refs:
            if not isinstance(ref, dict):
                continue
            extra = ref.get('extra')
            if not isinstance(extra, dict):
                continue
            x = extra.get('x')
            y = extra.get('y')
            if x is None or y is None:
                continue
            points.append(_format_point(x, y))
            if len(points) >= 3:
                return tuple(points)
    return tuple(points)


def _format_point(x: object, y: object) -> str:
    if isinstance(x, (int, float)) and isinstance(y, (int, float)):
        return f'({x:.1f}, {y:.1f})'
    return f'({x}, {y})'


def _format_rule_hits(issue_counts_by_rule: tuple[tuple[str, int], ...]) -> str:
    if not issue_counts_by_rule:
        return '无'
    return '；'.join(f'{rule_id}: {count}' for rule_id, count in issue_counts_by_rule)


def _format_points(points: tuple[str, ...]) -> str:
    if not points:
        return '无'
    return '、'.join(points)


def _format_report_refs(report_md_rel: str | None, report_docx_rel: str | None) -> str:
    parts: list[str] = []
    if report_md_rel:
        parts.append(f'Markdown `{report_md_rel}`')
    if report_docx_rel:
        parts.append(f'DOCX `{report_docx_rel}`')
    return '；'.join(parts) if parts else '无'


def _build_strict_note(ruleset_dir: Path | None) -> str | None:
    if ruleset_dir is None:
        return None
    loaded = load_ruleset_dir(ruleset_dir)
    for config in loaded.rule_configs:
        if config.rule_id != 'wire.floating_endpoints':
            continue
        if not config.enabled or config.severity != Severity.ERROR:
            return None
        scope = '全部图纸类型' if not config.applies_to else '、'.join(config.applies_to)
        return f'wire.floating_endpoints 已按 error 判定，适用图纸类型：{scope}。'
    return None


def _read_json(path: Path) -> dict | list:
    return json.loads(path.read_text(encoding='utf-8-sig'))


def _path_leaf(path_text: str) -> str:
    return PureWindowsPath(path_text).name or Path(path_text).name or path_text
