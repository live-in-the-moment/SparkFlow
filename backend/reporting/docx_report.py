from __future__ import annotations

import json
from collections.abc import Iterable
from pathlib import Path

from ..contracts import AuditReport, Issue
from .formal import build_formal_issue_details


def write_docx_report(report: AuditReport, out_path: Path) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    formal_issue_details = build_formal_issue_details(report)

    title = doc.add_paragraph("SparkFlow 审图报告")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"created_at: {report.created_at}")
    doc.add_paragraph(f"input_path: {report.input_path}")
    doc.add_paragraph(f"input_sha256: {report.input_sha256}")
    doc.add_paragraph(f"parser: {report.parser}")
    doc.add_paragraph(f"rule_version: {report.rule_version}")
    doc.add_paragraph(f"passed: {str(report.passed).lower()}")
    rule_hit_notes = report.summary.get('rule_hit_notes') if isinstance(report.summary, dict) else None

    if report.summary:
        doc.add_heading("Summary", level=1)
        _kv_table(doc, [(str(k), _short_json(v)) for k, v in report.summary.items() if k != 'rule_hit_notes'])

    if isinstance(rule_hit_notes, list) and rule_hit_notes:
        doc.add_heading("Rule Hit Notes", level=1)
        _rule_hit_note_table(doc, rule_hit_notes)

    if report.artifacts:
        doc.add_heading("Artifacts", level=1)
        _kv_table(doc, [(str(k), str(v)) for k, v in report.artifacts.items()])

    doc.add_heading("Issues", level=1)
    if not report.issues:
        doc.add_paragraph("无问题。")
    else:
        _issue_table(doc, formal_issue_details)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _kv_table(doc, rows: list[tuple[str, str]]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "字段"
    hdr[1].text = "值"
    for k, v in rows:
        r = table.add_row().cells
        r[0].text = k
        r[1].text = v


def _issue_table(doc, issues: Iterable) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=1, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "severity"
    hdr[2].text = "rule_id"
    hdr[3].text = "message"
    hdr[4].text = "article_clause_mapping"
    hdr[5].text = "risk_level"
    hdr[6].text = "confidence"
    hdr[7].text = "remediation"
    hdr[8].text = "refs"
    for i, detail in enumerate(list(issues), start=1):
        it = detail.issue
        r = table.add_row().cells
        r[0].text = str(i)
        r[1].text = str(it.severity.value)
        r[2].text = str(it.rule_id)
        r[3].text = str(it.message)
        r[4].text = str(detail.article_clause_mapping)
        r[5].text = str(detail.risk_level)
        r[6].text = str(detail.confidence)
        r[7].text = str(detail.remediation)
        ref_parts = []
        for ref in it.refs[:3]:
            piece = f"{ref.kind}:{ref.id}"
            extra = ref.extra or {}
            x = extra.get("x")
            y = extra.get("y")
            if x is not None and y is not None:
                piece = f"{piece} ({x},{y})"
            ref_parts.append(piece)
        r[8].text = " ".join(ref_parts)


def _rule_hit_note_table(doc, notes: list[dict[str, object]]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "severity"
    hdr[2].text = "rule_id"
    hdr[3].text = "drawing_type"
    hdr[4].text = "meaning"
    hdr[5].text = "grading_reason"
    for idx, note in enumerate(notes, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = str(note.get('severity', 'info'))
        row[2].text = str(note.get('rule_id', ''))
        row[3].text = str(note.get('drawing_type_label', ''))
        row[4].text = str(note.get('meaning', ''))
        row[5].text = str(note.get('grading_reason', ''))


def _short_json(v: object) -> str:
    try:
        s = json.dumps(v, ensure_ascii=False)
    except Exception:
        s = str(v)
    if len(s) > 1600:
        return s[:1600] + "..."
    return s
