from __future__ import annotations

import hashlib
import json
from collections import Counter
from dataclasses import dataclass
from pathlib import Path, PureWindowsPath

from ..contracts import AuditReport, Issue, ObjectRef, Severity
from .formal import build_formal_issue_details


@dataclass(frozen=True)
class RectificationChecklistArtifacts:
    markdown_path: Path
    docx_path: Path
    json_path: Path


@dataclass(frozen=True)
class RectificationCoordinate:
    x: float | str
    y: float | str


@dataclass(frozen=True)
class RectificationIssueEntry:
    sequence: int
    drawing_path: str
    rule_id: str
    severity: str
    message: str
    coordinates: tuple[RectificationCoordinate, ...]
    coordinate_labels: tuple[str, ...]
    article_clause_mapping: str
    rectification_suggestion: str
    risk_level: str
    confidence: str
    report_json_rel: str | None
    report_md_rel: str | None
    report_docx_rel: str | None


@dataclass(frozen=True)
class RectificationDrawingEntry:
    rel_path: str
    drawing_type_label: str
    issue_count: int
    issue_counts_by_rule: tuple[tuple[str, int], ...]
    report_json_rel: str | None
    report_md_rel: str | None
    report_docx_rel: str | None
    issues: tuple[RectificationIssueEntry, ...]


@dataclass(frozen=True)
class RectificationChecklistModel:
    title: str
    created_at: str
    dataset_dir: str
    dataset_label: str
    rule_version: str
    source_run_dir: str
    dataset_summary_rel: str
    dataset_selection_rel: str
    included_failed_drawings: int
    excluded_passed_drawings: int
    excluded_skipped_drawings: int
    total_issue_count: int
    issues_by_rule: tuple[tuple[str, int], ...]
    drawings: tuple[RectificationDrawingEntry, ...]


def write_rectification_checklist(
    run_dir: Path,
    *,
    out_md: Path | None = None,
    out_docx: Path | None = None,
    out_json: Path | None = None,
    title: str | None = None,
    dataset_label: str | None = None,
) -> RectificationChecklistArtifacts:
    run_dir = run_dir.resolve()
    model = build_rectification_checklist_model(
        run_dir,
        title=title,
        dataset_label=dataset_label,
    )

    markdown_path = (out_md or (run_dir / 'rectification_checklist.md')).resolve()
    docx_path = (out_docx or (run_dir / 'rectification_checklist.docx')).resolve()
    json_path = (out_json or (run_dir / 'rectification_checklist.json')).resolve()

    markdown_path.parent.mkdir(parents=True, exist_ok=True)
    markdown_path.write_text(render_rectification_checklist_markdown(model), encoding='utf-8')
    write_rectification_checklist_docx(model, docx_path)
    json_path.write_text(
        json.dumps(render_rectification_checklist_json(model), ensure_ascii=False, indent=2),
        encoding='utf-8',
    )
    return RectificationChecklistArtifacts(
        markdown_path=markdown_path,
        docx_path=docx_path,
        json_path=json_path,
    )


def build_rectification_checklist_model(
    run_dir: Path,
    *,
    title: str | None = None,
    dataset_label: str | None = None,
) -> RectificationChecklistModel:
    summary_path = run_dir / 'dataset_summary.json'
    selection_path = run_dir / 'dataset_selection.json'
    if not summary_path.exists():
        raise FileNotFoundError(str(summary_path))
    if not selection_path.exists():
        raise FileNotFoundError(str(selection_path))

    summary = _read_json(summary_path)
    selection_rows = _read_json(selection_path)
    if not isinstance(summary, dict):
        raise ValueError('dataset_summary.json 必须是对象。')
    if not isinstance(selection_rows, list):
        raise ValueError('dataset_selection.json 必须是数组。')

    resolved_dataset_label = dataset_label or _path_leaf(str(summary.get('dataset_dir') or run_dir.name))
    resolved_title = title or f'{resolved_dataset_label} 严格审图整改清单'
    failures = [
        row
        for row in selection_rows
        if isinstance(row, dict)
        and str(row.get('drawing_class') or '') == 'supported_electrical'
        and str(row.get('status') or '') == 'failed'
    ]
    passed_supported = [
        row
        for row in selection_rows
        if isinstance(row, dict)
        and str(row.get('drawing_class') or '') == 'supported_electrical'
        and str(row.get('status') or '') == 'passed'
    ]
    skipped_rows = [
        row
        for row in selection_rows
        if isinstance(row, dict) and str(row.get('status') or '') == 'skipped'
    ]

    drawings = tuple(_load_failed_drawing_entry(run_dir, row) for row in failures)
    total_issue_count = sum(entry.issue_count for entry in drawings)
    summary_total_issues = sum(
        int(count)
        for count in (summary.get('issues_by_rule') or {}).values()
        if isinstance(count, int | float | str)
    )
    if total_issue_count != summary_total_issues:
        raise ValueError(
            f'整改清单问题数与严格运行汇总不一致：提取到 {total_issue_count}，汇总为 {summary_total_issues}。'
        )

    summary_failed_count = int((summary.get('counts') or {}).get('failed', 0) or 0)
    if len(drawings) != summary_failed_count:
        raise ValueError(
            f'整改清单失败图纸数与严格运行汇总不一致：提取到 {len(drawings)}，汇总为 {summary_failed_count}。'
        )

    issue_counter: Counter[str] = Counter()
    for drawing in drawings:
        for rule_id, count in drawing.issue_counts_by_rule:
            issue_counter[rule_id] += count

    return RectificationChecklistModel(
        title=resolved_title,
        created_at=str(summary.get('created_at') or ''),
        dataset_dir=str(summary.get('dataset_dir') or ''),
        dataset_label=resolved_dataset_label,
        rule_version=str(summary.get('rule_version') or ''),
        source_run_dir=str(run_dir),
        dataset_summary_rel=summary_path.relative_to(run_dir).as_posix(),
        dataset_selection_rel=selection_path.relative_to(run_dir).as_posix(),
        included_failed_drawings=len(drawings),
        excluded_passed_drawings=len(passed_supported),
        excluded_skipped_drawings=len(skipped_rows),
        total_issue_count=total_issue_count,
        issues_by_rule=tuple(sorted(issue_counter.items(), key=lambda item: (-item[1], item[0]))),
        drawings=drawings,
    )


def render_rectification_checklist_markdown(model: RectificationChecklistModel) -> str:
    lines: list[str] = []
    lines.append(f'# {model.title}')
    lines.append('')
    lines.append('## 1. 清单范围')
    lines.append('')
    lines.append(f'- 来源运行目录：`{model.source_run_dir}`')
    lines.append(f'- 数据集目录：{model.dataset_dir}')
    lines.append(f'- 生成时间：{model.created_at}')
    lines.append(f'- 规则版本：{model.rule_version}')
    lines.append(f'- 数据集汇总：`{model.dataset_summary_rel}`')
    lines.append(f'- 选图清单：`{model.dataset_selection_rel}`')
    lines.append('- 仅纳入 `supported_electrical` 且 `status=failed` 的图纸，不重跑审图。')
    lines.append('')
    lines.append('## 2. 覆盖统计')
    lines.append('')
    lines.append(f'- 纳入整改清单的失败电气图纸：{model.included_failed_drawings}')
    lines.append(f'- 排除的已通过电气图纸：{model.excluded_passed_drawings}')
    lines.append(f'- 排除的跳过图纸：{model.excluded_skipped_drawings}')
    lines.append(f'- 整改项总数：{model.total_issue_count}')
    if model.issues_by_rule:
        lines.append(f"- 规则命中汇总：{'；'.join(f'{rule_id}: {count}' for rule_id, count in model.issues_by_rule)}")
    lines.append('')
    lines.append('## 3. 图纸覆盖')
    lines.append('')
    if not model.drawings:
        lines.append('无失败图纸。')
    else:
        for index, drawing in enumerate(model.drawings, start=1):
            lines.append(f'{index}. {drawing.rel_path}')
            lines.append(f'   - 图纸类型：{drawing.drawing_type_label}')
            lines.append(f'   - 问题数：{drawing.issue_count}')
            lines.append(f'   - 规则命中：{_format_rule_hits(drawing.issue_counts_by_rule)}')
            lines.append(f'   - 单图报告：{_format_report_refs(drawing.report_json_rel, drawing.report_md_rel, drawing.report_docx_rel)}')
    lines.append('')
    lines.append('## 4. 整改明细')
    lines.append('')
    if not model.drawings:
        lines.append('无。')
    else:
        for index, drawing in enumerate(model.drawings, start=1):
            lines.append(f'### 4.{index} {drawing.rel_path}')
            lines.append('')
            lines.append(f'- 图纸类型：{drawing.drawing_type_label}')
            lines.append(f'- 问题数：{drawing.issue_count}')
            lines.append(f'- 单图报告：{_format_report_refs(drawing.report_json_rel, drawing.report_md_rel, drawing.report_docx_rel)}')
            lines.append('')
            lines.append('| 序号 | 图纸路径 | rule_id | severity | message | 坐标 | 整改建议 | 单图报告 |')
            lines.append('| --- | --- | --- | --- | --- | --- | --- | --- |')
            for issue in drawing.issues:
                lines.append(
                    '| {sequence} | {drawing_path} | {rule_id} | {severity} | {message} | {coords} | {suggestion} | {report_refs} |'.format(
                        sequence=issue.sequence,
                        drawing_path=_escape_md_table(issue.drawing_path),
                        rule_id=_escape_md_table(issue.rule_id),
                        severity=_escape_md_table(issue.severity),
                        message=_escape_md_table(issue.message),
                        coords=_escape_md_table(_format_coordinates(issue.coordinate_labels)),
                        suggestion=_escape_md_table(issue.rectification_suggestion),
                        report_refs=_escape_md_table(
                            _format_report_refs(
                                issue.report_json_rel,
                                issue.report_md_rel,
                                issue.report_docx_rel,
                            )
                        ),
                    )
                )
            lines.append('')
    return '\n'.join(lines)


def render_rectification_checklist_json(model: RectificationChecklistModel) -> dict[str, object]:
    return {
        'title': model.title,
        'created_at': model.created_at,
        'dataset_dir': model.dataset_dir,
        'dataset_label': model.dataset_label,
        'rule_version': model.rule_version,
        'source_run_dir': model.source_run_dir,
        'source_artifacts': {
            'dataset_summary': model.dataset_summary_rel,
            'dataset_selection': model.dataset_selection_rel,
        },
        'counts': {
            'included_failed_drawings': model.included_failed_drawings,
            'excluded_passed_drawings': model.excluded_passed_drawings,
            'excluded_skipped_drawings': model.excluded_skipped_drawings,
            'total_issue_count': model.total_issue_count,
        },
        'issues_by_rule': [
            {'rule_id': rule_id, 'count': count}
            for rule_id, count in model.issues_by_rule
        ],
        'drawings': [
            {
                'rel_path': drawing.rel_path,
                'drawing_type_label': drawing.drawing_type_label,
                'issue_count': drawing.issue_count,
                'issue_counts_by_rule': [
                    {'rule_id': rule_id, 'count': count}
                    for rule_id, count in drawing.issue_counts_by_rule
                ],
                'report_references': {
                    'json': drawing.report_json_rel,
                    'markdown': drawing.report_md_rel,
                    'docx': drawing.report_docx_rel,
                },
                'issues': [
                    {
                        'sequence': issue.sequence,
                        'drawing_path': issue.drawing_path,
                        'rule_id': issue.rule_id,
                        'severity': issue.severity,
                        'message': issue.message,
                        'coordinates': [
                            {'x': coordinate.x, 'y': coordinate.y}
                            for coordinate in issue.coordinates
                        ],
                        'coordinate_labels': list(issue.coordinate_labels),
                        'article_clause_mapping': issue.article_clause_mapping,
                        'rectification_suggestion': issue.rectification_suggestion,
                        'risk_level': issue.risk_level,
                        'confidence': issue.confidence,
                        'report_references': {
                            'json': issue.report_json_rel,
                            'markdown': issue.report_md_rel,
                            'docx': issue.report_docx_rel,
                        },
                    }
                    for issue in drawing.issues
                ],
            }
            for drawing in model.drawings
        ],
    }


def write_rectification_checklist_docx(model: RectificationChecklistModel, out_path: Path) -> Path:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    _configure_docx_styles(doc, qn, Pt)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = doc.styles['Title']
    title.add_run(model.title)

    doc.add_heading('1. 清单范围', level=1)
    _doc_line(doc, f'来源运行目录：{model.source_run_dir}')
    _doc_line(doc, f'数据集目录：{model.dataset_dir}')
    _doc_line(doc, f'生成时间：{model.created_at}')
    _doc_line(doc, f'规则版本：{model.rule_version}')
    _doc_line(doc, f'数据集汇总：{model.dataset_summary_rel}')
    _doc_line(doc, f'选图清单：{model.dataset_selection_rel}')
    _doc_line(doc, '仅纳入 supported_electrical 且 status=failed 的图纸，不重跑审图。')

    doc.add_heading('2. 覆盖统计', level=1)
    _doc_line(doc, f'纳入整改清单的失败电气图纸：{model.included_failed_drawings}')
    _doc_line(doc, f'排除的已通过电气图纸：{model.excluded_passed_drawings}')
    _doc_line(doc, f'排除的跳过图纸：{model.excluded_skipped_drawings}')
    _doc_line(doc, f'整改项总数：{model.total_issue_count}')
    if model.issues_by_rule:
        _doc_line(doc, f"规则命中汇总：{'；'.join(f'{rule_id}: {count}' for rule_id, count in model.issues_by_rule)}")

    doc.add_heading('3. 图纸覆盖', level=1)
    if not model.drawings:
        _doc_line(doc, '无失败图纸。')
    else:
        for index, drawing in enumerate(model.drawings, start=1):
            doc.add_heading(f'{index}. {drawing.rel_path}', level=2)
            _doc_line(doc, f'图纸类型：{drawing.drawing_type_label}')
            _doc_line(doc, f'问题数：{drawing.issue_count}')
            _doc_line(doc, f'规则命中：{_format_rule_hits(drawing.issue_counts_by_rule)}')
            _doc_line(doc, f'单图报告：{_format_report_refs(drawing.report_json_rel, drawing.report_md_rel, drawing.report_docx_rel)}')

    doc.add_heading('4. 整改明细', level=1)
    if not model.drawings:
        _doc_line(doc, '无。')
    else:
        for index, drawing in enumerate(model.drawings, start=1):
            doc.add_heading(f'4.{index} {drawing.rel_path}', level=2)
            _doc_line(doc, f'图纸类型：{drawing.drawing_type_label}')
            _doc_line(doc, f'问题数：{drawing.issue_count}')
            _doc_line(doc, f'单图报告：{_format_report_refs(drawing.report_json_rel, drawing.report_md_rel, drawing.report_docx_rel)}')

            table = doc.add_table(rows=1, cols=8)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = '序号'
            hdr[1].text = '图纸路径'
            hdr[2].text = 'rule_id'
            hdr[3].text = 'severity'
            hdr[4].text = 'message'
            hdr[5].text = '坐标'
            hdr[6].text = '整改建议'
            hdr[7].text = '单图报告'
            for issue in drawing.issues:
                row = table.add_row().cells
                row[0].text = str(issue.sequence)
                row[1].text = issue.drawing_path
                row[2].text = issue.rule_id
                row[3].text = issue.severity
                row[4].text = issue.message
                row[5].text = _format_coordinates(issue.coordinate_labels)
                row[6].text = issue.rectification_suggestion
                row[7].text = _format_report_refs(issue.report_json_rel, issue.report_md_rel, issue.report_docx_rel)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _load_failed_drawing_entry(run_dir: Path, row: dict[str, object]) -> RectificationDrawingEntry:
    rel_path = str(row.get('rel_path') or '')
    file_dir = _dataset_file_out_dir(run_dir, rel_path)
    report_path = file_dir / 'report.json'
    if not report_path.exists():
        raise FileNotFoundError(str(report_path))
    report_payload = _read_json(report_path)
    if not isinstance(report_payload, dict):
        raise ValueError(f'{report_path} 必须是对象。')
    report = _deserialize_report(report_payload)
    details = build_formal_issue_details(report)
    classification = report.summary.get('classification') if isinstance(report.summary, dict) else {}
    drawing_type_label = str(classification.get('drawing_type_label') or '未标注')
    report_md_path = file_dir / 'report.md'
    report_docx_path = file_dir / 'report.docx'

    issues: list[RectificationIssueEntry] = []
    issue_counts = Counter(issue.rule_id for issue in report.issues)
    for sequence, detail in enumerate(details, start=1):
        coordinates = _extract_coordinates(detail.issue)
        coordinate_labels = tuple(_format_point(coordinate.x, coordinate.y) for coordinate in coordinates)
        issues.append(
            RectificationIssueEntry(
                sequence=sequence,
                drawing_path=rel_path,
                rule_id=detail.issue.rule_id,
                severity=detail.issue.severity.value,
                message=detail.issue.message,
                coordinates=coordinates,
                coordinate_labels=coordinate_labels,
                article_clause_mapping=detail.article_clause_mapping,
                rectification_suggestion=_build_rectification_suggestion(detail.remediation, coordinate_labels),
                risk_level=detail.risk_level,
                confidence=detail.confidence,
                report_json_rel=report_path.relative_to(run_dir).as_posix(),
                report_md_rel=report_md_path.relative_to(run_dir).as_posix() if report_md_path.exists() else None,
                report_docx_rel=report_docx_path.relative_to(run_dir).as_posix() if report_docx_path.exists() else None,
            )
        )

    return RectificationDrawingEntry(
        rel_path=rel_path,
        drawing_type_label=drawing_type_label,
        issue_count=len(issues),
        issue_counts_by_rule=tuple(sorted(issue_counts.items(), key=lambda item: (-item[1], item[0]))),
        report_json_rel=report_path.relative_to(run_dir).as_posix(),
        report_md_rel=report_md_path.relative_to(run_dir).as_posix() if report_md_path.exists() else None,
        report_docx_rel=report_docx_path.relative_to(run_dir).as_posix() if report_docx_path.exists() else None,
        issues=tuple(issues),
    )


def _deserialize_report(payload: dict[str, object]) -> AuditReport:
    issues_payload = payload.get('issues')
    issues: list[Issue] = []
    if isinstance(issues_payload, list):
        for item in issues_payload:
            if not isinstance(item, dict):
                continue
            refs_payload = item.get('refs')
            refs: list[ObjectRef] = []
            if isinstance(refs_payload, list):
                for ref in refs_payload:
                    if not isinstance(ref, dict):
                        continue
                    source_entity_ids = ref.get('source_entity_ids')
                    refs.append(
                        ObjectRef(
                            kind=str(ref.get('kind') or ''),
                            id=str(ref.get('id') or ''),
                            source_entity_ids=tuple(
                                str(entity_id)
                                for entity_id in (source_entity_ids if isinstance(source_entity_ids, list) else [])
                            ),
                            extra=dict(ref.get('extra') or {}) if isinstance(ref.get('extra'), dict) else {},
                        )
                    )
            issues.append(
                Issue(
                    rule_id=str(item.get('rule_id') or ''),
                    severity=Severity(str(item.get('severity') or 'info')),
                    message=str(item.get('message') or ''),
                    refs=tuple(refs),
                )
            )

    return AuditReport(
        created_at=str(payload.get('created_at') or ''),
        input_path=str(payload.get('input_path') or ''),
        input_sha256=str(payload.get('input_sha256') or ''),
        parser=str(payload.get('parser') or ''),
        rule_version=str(payload.get('rule_version') or ''),
        summary=dict(payload.get('summary') or {}) if isinstance(payload.get('summary'), dict) else {},
        artifacts=dict(payload.get('artifacts') or {}) if isinstance(payload.get('artifacts'), dict) else {},
        issues=tuple(issues),
    )


def _extract_coordinates(issue: Issue) -> tuple[RectificationCoordinate, ...]:
    points: list[RectificationCoordinate] = []
    for ref in issue.refs:
        extra = ref.extra or {}
        x = extra.get('x')
        y = extra.get('y')
        if x is None or y is None:
            continue
        points.append(RectificationCoordinate(x=x, y=y))
    return tuple(points)


def _build_rectification_suggestion(remediation: str, coordinate_labels: tuple[str, ...]) -> str:
    if coordinate_labels:
        return f"优先复核坐标 {'、'.join(coordinate_labels)} 附近对象；{remediation}"
    return remediation


def _format_coordinates(coordinate_labels: tuple[str, ...]) -> str:
    if not coordinate_labels:
        return '无'
    return '；'.join(coordinate_labels)


def _format_point(x: float | str, y: float | str) -> str:
    if isinstance(x, (int, float)) and isinstance(y, (int, float)):
        return f'({x:.3f}, {y:.3f})'
    return f'({x}, {y})'


def _format_rule_hits(issue_counts_by_rule: tuple[tuple[str, int], ...]) -> str:
    if not issue_counts_by_rule:
        return '无'
    return '；'.join(f'{rule_id}: {count}' for rule_id, count in issue_counts_by_rule)


def _format_report_refs(
    report_json_rel: str | None,
    report_md_rel: str | None,
    report_docx_rel: str | None,
) -> str:
    parts: list[str] = []
    if report_json_rel:
        parts.append(f'JSON `{report_json_rel}`')
    if report_md_rel:
        parts.append(f'Markdown `{report_md_rel}`')
    if report_docx_rel:
        parts.append(f'DOCX `{report_docx_rel}`')
    return '；'.join(parts) if parts else '无'


def _escape_md_table(value: str) -> str:
    return value.replace('|', '\\|').replace('\n', '<br>')


def _configure_docx_styles(doc, qn, Pt) -> None:
    for style_name, font_name, size in (
        ('Normal', '宋体', 10.5),
        ('Title', '微软雅黑', 18),
        ('Heading 1', '微软雅黑', 14),
        ('Heading 2', '微软雅黑', 12),
    ):
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), font_name)


def _doc_line(doc, text: str) -> None:
    doc.add_paragraph(text)


def _dataset_file_out_dir(run_dir: Path, rel_path: str) -> Path:
    path = Path(*rel_path.split('/'))
    token = hashlib.sha1(rel_path.encode('utf-8')).hexdigest()[:8]
    return run_dir / 'files' / path.parent / f'{path.stem}__{token}'


def _read_json(path: Path) -> dict | list:
    return json.loads(path.read_text(encoding='utf-8-sig'))


def _path_leaf(path_text: str) -> str:
    return PureWindowsPath(path_text).name or Path(path_text).name or path_text
